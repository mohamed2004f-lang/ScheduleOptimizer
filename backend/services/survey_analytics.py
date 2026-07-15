"""تحليل وتصدير تقارير الاستبيانات (Excel متعدد الأوراق)."""

from __future__ import annotations

import datetime
import io
import math
import tempfile
from typing import Any

import pandas as pd
from flask import send_file

from backend.core.survey_platform import (
    LINK_TYPE_LABELS_AR,
    RESPONDENT_ROLE_LABELS,
    survey_template_intro,
)
from backend.database.database import fetch_table_columns, schedule_pk_column, table_exists
from backend.services.accreditation_metrics import suggest_compliance_status
from backend.services.multi_surveys import aggregate_template, get_template_by_code, list_templates
from backend.services.quality_metrics import _avg_eval_score, term_label_from_conn
from backend.core.arabic_export import (
    docx_add_rtl_heading,
    docx_add_rtl_paragraph,
    docx_fill_rtl_table,
    excel_arabic_workbook_formats,
    pdf_arabic_extra_css,
    set_docx_paragraph_rtl,
    set_docx_run_arabic_font,
    write_excel_sheet_rtl,
)
from backend.services.utilities import excel_response_from_frames, schedule_semester_matches_current_term

COMPLIANCE_STATUS_LABELS_AR: dict[str, str] = {
    "met": "متحقق",
    "partial": "جزئي",
    "gap": "فجوة",
    "in_progress": "قيد الإنجاز",
}

SCORE_CLASS_LABELS_AR: dict[str, str] = {
    "excellent": "ممتاز",
    "good": "جيد",
    "needs_improvement": "يحتاج تحسين",
    "critical": "حرج",
    "pending": "بانتظار التجميع",
}

# حد التجميع لتقييم المقرر: نسبة من المسجّلين (افتراضي 50% — قابل للتعديل من لوحة القيادة)
COURSE_EVAL_RESPONSE_RATE_SETTING_KEY = "course_eval_response_rate_percent"
COURSE_EVAL_DEFAULT_RATE_PERCENT = 50
COURSE_EVAL_RATE_MIN_PERCENT = 5
COURSE_EVAL_RATE_MAX_PERCENT = 100
# عند غياب بيانات التسجيل لا تُطبَّق النسبة — لا يُعرض التجميع
COURSE_EVAL_NO_ENROLLMENT_MIN = 2**30

COLLEGE_NAME_AR = "كلية الهندسة"


def _read_system_setting(cur, key: str, default: str = "") -> str:
    try:
        row = cur.execute("SELECT value FROM system_settings WHERE key = ?", (key,)).fetchone()
        return (row[0] or default) if row else default
    except Exception:
        return default


def get_course_eval_response_rate_percent(conn=None) -> int:
    """نسبة الاستجابة المطلوبة لإظهار نتيجة تقييم المقرر (5–100، افتراضي 50)."""
    def _read(c):
        raw = _read_system_setting(
            c.cursor(),
            COURSE_EVAL_RESPONSE_RATE_SETTING_KEY,
            str(COURSE_EVAL_DEFAULT_RATE_PERCENT),
        )
        try:
            pct = int(str(raw).strip())
        except (TypeError, ValueError):
            pct = COURSE_EVAL_DEFAULT_RATE_PERCENT
        return max(COURSE_EVAL_RATE_MIN_PERCENT, min(COURSE_EVAL_RATE_MAX_PERCENT, pct))

    if conn is not None:
        return _read(conn)
    from backend.services.utilities import get_connection

    with get_connection() as c:
        return _read(c)


def get_course_eval_response_rate(conn=None) -> float:
    return get_course_eval_response_rate_percent(conn) / 100.0


def set_course_eval_response_rate_percent(conn, percent: int) -> int:
    """يحفظ نسبة تجميع تقييم المقرر ويعيد القيمة المطبّقة."""
    pct = max(
        COURSE_EVAL_RATE_MIN_PERCENT,
        min(COURSE_EVAL_RATE_MAX_PERCENT, int(percent)),
    )
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO system_settings (key, value) VALUES (?, ?)
        ON CONFLICT (key) DO UPDATE SET value = EXCLUDED.value
        """,
        (COURSE_EVAL_RESPONSE_RATE_SETTING_KEY, str(pct)),
    )
    conn.commit()
    return pct


def format_course_eval_aggregation_policy(conn=None) -> str:
    return f"{get_course_eval_response_rate_percent(conn)}% من المسجّلين"

RECOMMENDATION_BY_CLASS: dict[str, str] = {
    "excellent": "الحفاظ على الممارسة وتوثيقها كمثال يُحتذى به.",
    "good": "مراقبة دورية ومتابعة في الاجتماعات الفصلية.",
    "needs_improvement": "وضع خطة تحسين فصلية مع مسؤول ومؤشر متابعة.",
    "critical": "اجتماع عاجل مع المعنيين وإجراء تصحيحي خلال 30 يوماً.",
    "pending": "—",
}


def classify_item_score(percent: float | None) -> str:
    if percent is None:
        return "pending"
    p = float(percent)
    if p >= 80:
        return "excellent"
    if p >= 70:
        return "good"
    if p >= 50:
        return "needs_improvement"
    return "critical"


def classify_section_score(percent: float | None) -> str:
    """تصنيف نتيجة الشعبة في تقارير الشواهد."""
    if percent is None:
        return "pending"
    p = float(percent)
    if p >= 90:
        return "excellent"
    if p >= 75:
        return "very_good"
    if p >= 65:
        return "good"
    return "needs_improvement"


SECTION_SCORE_LABELS_AR: dict[str, str] = {
    "excellent": "ممتاز (≥90%)",
    "very_good": "جيد جداً (75–89%)",
    "good": "جيد (65–74.99%)",
    "needs_improvement": "يحتاج تحسين (<65%)",
    "pending": "ناقص التجميع",
}

SECTION_SCORE_BUCKET_ORDER: tuple[str, ...] = (
    "excellent",
    "very_good",
    "good",
    "needs_improvement",
    "pending",
)


def classify_compliance_status(score_percent: float | None) -> str:
    return suggest_compliance_status(score_percent)


def accreditation_links_for(
    template_code: str,
    conn=None,
    *,
    semester: str | None = None,
    department_id: int | None = None,
    catalog_version: str | None = None,
) -> list[dict[str, Any]]:
    """روابط الاعتماد المعروضة — من bindings وقواعد الكتالوج فقط."""
    if conn is None:
        return []
    from backend.services.survey_accreditation_links import (
        resolve_survey_accreditation_links,
    )

    return resolve_survey_accreditation_links(
        conn,
        template_code,
        catalog_version=catalog_version,
        semester=semester,
        department_id=department_id,
    )


def _department_label(conn, department_id: int | None) -> str:
    if department_id is None:
        return "الكلية (كل الأقسام)"
    cur = conn.cursor()
    row = cur.execute(
        "SELECT COALESCE(name_ar, code, '') AS label FROM departments WHERE id = ? LIMIT 1",
        (int(department_id),),
    ).fetchone()
    if not row:
        return f"قسم #{department_id}"
    name = row[0] if hasattr(row, "__getitem__") else ""
    if not name and hasattr(row, "get"):
        name = row.get("label") or ""
    return str(name or "").strip() or f"قسم #{department_id}"


def _enrich_questions(questions: list[dict]) -> list[dict]:
    out: list[dict] = []
    for q in questions:
        pct = q.get("score_percent")
        cls = classify_item_score(pct if pct is not None else None)
        out.append(
            {
                **q,
                "classification": cls,
                "classification_ar": SCORE_CLASS_LABELS_AR.get(cls, cls),
                "recommendation_ar": RECOMMENDATION_BY_CLASS.get(cls, "—"),
            }
        )
    return out


def _weakest_strongest(questions: list[dict]) -> tuple[str, str]:
    scored = [q for q in questions if q.get("score_percent") is not None]
    if not scored:
        return "—", "—"
    weakest = min(scored, key=lambda x: float(x["score_percent"]))
    strongest = max(scored, key=lambda x: float(x["score_percent"]))
    return (weakest.get("label_ar") or "—")[:80], (strongest.get("label_ar") or "—")[:80]


def _primary_accreditation_label(
    conn,
    template_code: str,
    *,
    semester: str | None = None,
    department_id: int | None = None,
) -> str:
    links = accreditation_links_for(
        template_code, conn, semester=semester, department_id=department_id
    )
    if not links:
        return "—"
    first = links[0]
    return f"{first.get('indicator_code', '')} — {first.get('indicator_title_ar', '')}"


def build_survey_report(
    conn,
    template_code: str,
    *,
    semester: str | None = None,
    department_id: int | None = None,
) -> dict[str, Any]:
    """تقرير تحليلي لاستبيان واحد."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    agg = aggregate_template(conn, template_code, semester=sem, department_id=department_id)
    tpl = get_template_by_code(conn, template_code) or {}
    questions = _enrich_questions(agg.get("questions") or [])
    weakest, strongest = _weakest_strongest(questions)
    score = agg.get("overall_score_percent")
    compliance = classify_compliance_status(score if agg.get("aggregated") else None)
    intro = survey_template_intro(template_code)

    return {
        **agg,
        "template_code": template_code,
        "semester": sem,
        "department_id": department_id,
        "department_label": _department_label(conn, department_id),
        "respondent_role": tpl.get("respondent_role"),
        "respondent_label": RESPONDENT_ROLE_LABELS.get(
            (tpl.get("respondent_role") or "").strip(), "—"
        ),
        "scope_note_ar": intro.get("scope_note_ar") or "",
        "questions": questions,
        "weakest_item": weakest,
        "strongest_item": strongest,
        "compliance_status": compliance,
        "compliance_status_ar": COMPLIANCE_STATUS_LABELS_AR.get(compliance, compliance),
        "accreditation_links": accreditation_links_for(
            template_code, conn, semester=sem, department_id=department_id
        ),
        "primary_accreditation": _primary_accreditation_label(
            conn, template_code, semester=sem, department_id=department_id
        ),
        "recommendations": generate_recommendations(questions, agg.get("title_ar") or template_code),
    }


def generate_recommendations(questions: list[dict], title_ar: str = "") -> list[str]:
    """توصيات آلية من البنود الأضعف."""
    critical = [q for q in questions if q.get("classification") == "critical"]
    weak = [q for q in questions if q.get("classification") == "needs_improvement"]
    recs: list[str] = []
    label = (title_ar or "الاستبيان").strip()
    for q in critical[:3]:
        recs.append(
            f"[حرج] {label}: معالجة عاجلة لبند «{(q.get('label_ar') or '')[:60]}» "
            f"({q.get('score_percent')}%)."
        )
    for q in weak[:3]:
        recs.append(
            f"[تحسين] {label}: خطة فصلية لبند «{(q.get('label_ar') or '')[:60]}» "
            f"({q.get('score_percent')}%)."
        )
    if not recs and questions:
        scored = [q for q in questions if q.get("score_percent") is not None]
        if scored and all(q.get("classification") in ("excellent", "good") for q in scored):
            recs.append(f"نتائج {label} إيجابية — يُوصى بتوثيق الممارسات الجيدة كشواهد اعتماد.")
    if not recs:
        recs.append("لا توجد توصيات — بانتظار اكتمال التجميع أو إدخال إجابات كافية.")
    return recs


def _legacy_eval_avg_expr() -> str:
    return """(
        COALESCE(e.instructor_punctuality, 0) +
        COALESCE(e.course_clarity, 0) +
        COALESCE(e.assessment_fairness, 0) +
        COALESCE(e.material_relevance, 0) +
        COALESCE(e.communication_quality, 0)
    ) / 5.0"""


def _eval_rating_subquery() -> str:
    legacy = _legacy_eval_avg_expr()
    return f"""COALESCE(
        (SELECT AVG(a.rating * 1.0)
         FROM evaluation_survey_answers a
         WHERE a.evaluation_id = e.id),
        {legacy}
    )"""


def _course_eval_dept_filter(conn, department_id: int | None) -> tuple[str, list[Any]]:
    if department_id is None:
        return "", []
    pk = schedule_pk_column(conn)
    return (
        f"""
        AND EXISTS (
            SELECT 1 FROM schedule sch
            WHERE sch.{pk} = e.section_id AND sch.department_id = ?
        )
        """,
        [int(department_id)],
    )


def course_eval_min_required(enrolled: int, *, conn=None, response_count: int = 0) -> int:
    """
    الحد الأدنى لإظهار نتيجة تقييم المقرر:
    - نسبة قابلة للإعداد من المسجّلين (افتراضي 50%، مقرب للأعلى)
    - عند غياب بيانات التسجيل: لا يُعرض التجميع
    """
    del response_count  # متوافق مع استدعاءات قديمة
    if enrolled > 0:
        rate = get_course_eval_response_rate(conn)
        return max(1, int(math.ceil(enrolled * rate)))
    return COURSE_EVAL_NO_ENROLLMENT_MIN


def course_eval_is_aggregated(response_count: int, enrolled: int, *, conn=None) -> bool:
    min_req = course_eval_min_required(enrolled, conn=conn)
    return int(response_count) >= min_req


def _course_registration_count(conn, course_name: str) -> int:
    cname = (course_name or "").strip()
    if not cname or not table_exists(conn, "registrations"):
        return 0
    cur = conn.cursor()
    row = cur.execute(
        """
        SELECT COUNT(DISTINCT student_id)
        FROM registrations
        WHERE lower(trim(course_name)) = lower(trim(?))
        """,
        (cname,),
    ).fetchone()
    return int(_row_val(row, 0) or 0)


def _course_section_count(conn, course_name: str, semester: str) -> int:
    """عدد صفوف الشعب في الجدول (قد يتضمن محاضرات متعددة لنفس الشعبة المنطقية)."""
    cname = (course_name or "").strip()
    if not cname or not table_exists(conn, "schedule"):
        return 1
    cur = conn.cursor()
    pk = schedule_pk_column(conn)
    sem = (semester or "").strip()
    row = cur.execute(
        f"""
        SELECT COUNT(DISTINCT sch.{pk})
        FROM schedule sch
        WHERE lower(trim(sch.course_name)) = lower(trim(?))
          AND (COALESCE(sch.semester, '') = ? OR ? = '')
        """,
        (cname, sem, sem),
    ).fetchone()
    return max(1, int(_row_val(row, 0) or 0))


def _course_teaching_group_count(
    conn,
    course_name: str,
    semester: str,
    *,
    instructor_id: int | None = None,
    department_id: int | None = None,
) -> int:
    """
    عدد مجموعات التدريس في الفصل — من جدول teaching_groups عند توفره، وإلا تقدير قديم.
    """
    from backend.services import teaching_groups as tg_svc

    cname = (course_name or "").strip().lower()
    if not cname:
        return 1
    if tg_svc.semester_has_teaching_groups(conn, semester):
        groups = tg_svc.list_teaching_groups(
            conn,
            semester=semester,
            department_id=department_id,
            course_name=(course_name or "").strip(),
        )
        if instructor_id is not None:
            groups = [g for g in groups if int(g.get("instructor_id") or 0) == int(instructor_id)]
        if groups:
            return max(1, len(groups))
    legacy_groups: set[tuple[str, int]] = set()
    for sec in _list_schedule_sections_for_term(conn, semester, department_id=department_id):
        if (sec.get("course_name") or "").strip().lower() != cname:
            continue
        iid = int(sec.get("instructor_id") or 0)
        if instructor_id is not None and iid != int(instructor_id):
            continue
        if iid > 0:
            legacy_groups.add((cname, iid))
        else:
            legacy_groups.add((cname, -int(sec.get("section_id") or 0)))
    return max(1, len(legacy_groups))


def _college_course_eval_enrolled(
    conn,
    semester: str,
    department_id: int | None = None,
) -> int:
    """عدد الطلاب المسجّلين في المقررات التي وُجد لها تقييم في الفصل."""
    if not table_exists(conn, "registrations") or not table_exists(conn, "course_evaluations"):
        return 0
    cur = conn.cursor()
    if department_id is not None:
        pk = schedule_pk_column(conn)
        row = cur.execute(
            f"""
            SELECT COUNT(DISTINCT r.student_id)
            FROM registrations r
            INNER JOIN course_evaluations e
                ON lower(trim(e.course_name)) = lower(trim(r.course_name))
               AND e.semester = ?
            WHERE EXISTS (
                SELECT 1 FROM schedule sch
                WHERE sch.{pk} = e.section_id AND sch.department_id = ?
            )
            """,
            (semester, int(department_id)),
        ).fetchone()
    else:
        row = cur.execute(
            """
            SELECT COUNT(DISTINCT r.student_id)
            FROM registrations r
            INNER JOIN course_evaluations e
                ON lower(trim(e.course_name)) = lower(trim(r.course_name))
               AND e.semester = ?
            """,
            (semester,),
        ).fetchone()
    return int(_row_val(row, 0) or 0)


def section_enrolled_count(
    conn,
    course_name: str,
    semester: str,
    *,
    section_count: int = 1,
    teaching_group_id: int | None = None,
) -> int:
    """
    عدد مسجّلي الشعبة/المجموعة.
    إن وُجد teaching_group_id يُستخدم COUNT الفعلي؛ وإلا تقدير قديم ÷ عدد الشعب.
    """
    if teaching_group_id is not None and int(teaching_group_id) > 0:
        from backend.services import teaching_groups as tg_svc
        return tg_svc.count_registrations_for_teaching_group(conn, int(teaching_group_id))
    total = _course_registration_count(conn, course_name)
    if total <= 0:
        return 0
    n_sections = max(1, int(section_count or 1))
    return max(1, int(math.ceil(total / n_sections)))


def _row_val(row, idx: int = 0, key: str | None = None):
    if row is None:
        return None
    if key and hasattr(row, "keys"):
        try:
            return row[key]
        except (KeyError, TypeError):
            pass
    try:
        return row[idx]
    except (IndexError, TypeError):
        return None


def _fetch_course_eval_section_groups(
    conn,
    semester: str,
    department_id: int | None = None,
) -> list[dict[str, Any]]:
    """وحدات التقييم — مجموعة تدريس أو شعبة (section_id + course + instructor)."""
    if not table_exists(conn, "course_evaluations"):
        return []
    cur = conn.cursor()
    pk = schedule_pk_column(conn)
    dept_sql, dept_params = _course_eval_dept_filter(conn, department_id)
    ce_cols = {c.lower() for c in fetch_table_columns(conn, "course_evaluations")}
    has_tg = "teaching_group_id" in ce_cols and table_exists(conn, "teaching_groups")
    tg_select = ""
    tg_join = ""
    if has_tg:
        tg_select = ", COALESCE(MAX(e.teaching_group_id), 0) AS teaching_group_id, COALESCE(MAX(tg.group_code), '') AS group_code"
        tg_join = " LEFT JOIN teaching_groups tg ON tg.id = e.teaching_group_id "
    rows = cur.execute(
        f"""
        SELECT COALESCE(MAX(e.section_id), 0) AS section_id,
               COALESCE(MAX(e.course_name), '') AS course_name,
               COALESCE(MAX(e.instructor_id), 0) AS instructor_id,
               COUNT(*) AS response_count,
               COALESCE(MAX(i.name), '') AS instructor_name,
               COALESCE(MAX(COALESCE(td.name_ar, td.code, d.name_ar, d.code)), '') AS department_name
               {tg_select}
        FROM course_evaluations e
        LEFT JOIN schedule sch ON sch.{pk} = e.section_id
        LEFT JOIN departments d ON d.id = sch.department_id
        LEFT JOIN instructors i ON i.id = e.instructor_id
        {tg_join}
        LEFT JOIN departments td ON td.id = tg.department_id
        WHERE e.semester = ? {dept_sql}
        GROUP BY COALESCE(NULLIF(e.teaching_group_id, 0), e.section_id), e.course_name, e.instructor_id
        ORDER BY COALESCE(MAX(e.course_name), ''), COALESCE(MAX(i.name), ''), COALESCE(MAX(e.section_id), 0)
        """,
        tuple([semester] + dept_params),
    ).fetchall()
    out: list[dict[str, Any]] = []
    seen_tg: set[int] = set()
    for r in rows:
        if hasattr(r, "keys"):
            d = dict(r)
        else:
            d = {
                "section_id": r[0],
                "course_name": r[1],
                "instructor_id": r[2],
                "response_count": r[3],
                "instructor_name": r[4],
                "department_name": r[5],
            }
            if has_tg and len(r) > 6:
                d["teaching_group_id"] = r[6]
                d["group_code"] = r[7]
        sid = int(d.get("section_id") or 0)
        tgid = int(d.get("teaching_group_id") or 0)
        if tgid > 0:
            if tgid in seen_tg:
                continue
            seen_tg.add(tgid)
            if not sid:
                from backend.services import teaching_groups as tg_svc

                sid = tg_svc.primary_section_id_for_group(conn, tgid)
        if not sid and not tgid:
            continue
        group_label = ""
        if tgid > 0:
            from backend.services import teaching_groups as tg_svc

            group_label = tg_svc.group_code_label(d.get("group_code"))
        out.append(
            {
                "section_id": sid,
                "teaching_group_id": tgid or None,
                "course_name": (d.get("course_name") or "").strip(),
                "instructor_id": int(d.get("instructor_id") or 0),
                "instructor_name": (d.get("instructor_name") or "").strip() or "—",
                "department_name": (d.get("department_name") or "").strip() or "—",
                "group_code_label": group_label or None,
                "response_count": int(d.get("response_count") or 0),
            }
        )
    return out


def _evaluated_section_ids(
    conn,
    semester: str,
    *,
    department_id: int | None = None,
) -> set[int]:
    """معرّفات الشعب التي وُجد لها تقييم واحد على الأقل في الفصل."""
    if not table_exists(conn, "course_evaluations"):
        return set()
    cur = conn.cursor()
    pk = schedule_pk_column(conn)
    dept_sql = ""
    params: list[Any] = [semester]
    if department_id is not None:
        dept_sql = f"""
            AND EXISTS (
                SELECT 1 FROM schedule sch
                WHERE sch.{pk} = e.section_id AND sch.department_id = ?
            )
        """
        params.append(int(department_id))
    rows = cur.execute(
        f"""
        SELECT DISTINCT e.section_id
        FROM course_evaluations e
        WHERE e.semester = ? {dept_sql}
        """,
        tuple(params),
    ).fetchall()
    out: set[int] = set()
    for row in rows:
        sid = int(_row_val(row, 0, "section_id") or 0)
        if sid:
            out.add(sid)
    return out


def _list_schedule_sections_for_term(
    conn,
    semester: str,
    *,
    department_id: int | None = None,
) -> list[dict[str, Any]]:
    """شعب الجدول المرتبطة بالفصل الدراسي (مع فلتر القسم اختيارياً)."""
    if not table_exists(conn, "schedule"):
        return []
    from backend.services.course_evaluations import (
        _instructor_id_by_name_map,
        _resolve_schedule_instructor_id,
        _schedule_section_id_expr,
    )

    sem = (semester or "").strip()
    pk = schedule_pk_column(conn)
    sid_expr = _schedule_section_id_expr(conn)
    cur = conn.cursor()
    name_map = _instructor_id_by_name_map(cur)
    dept_sql = ""
    dept_params: list[Any] = []
    if department_id is not None:
        dept_sql = " AND sch.department_id = ? "
        dept_params = [int(department_id)]

    rows = cur.execute(
        f"""
        SELECT {sid_expr} AS section_id,
               COALESCE(MAX(sch.course_name), '') AS course_name,
               COALESCE(MAX(sch.instructor_id), 0) AS instructor_id,
               COALESCE(MAX(sch.instructor), '') AS instructor_name,
               COALESCE(MAX(sch.semester), '') AS schedule_semester,
               COALESCE(MAX(d.name_ar), MAX(d.code), '') AS department_name
        FROM schedule sch
        LEFT JOIN departments d ON d.id = sch.department_id
        WHERE 1=1 {dept_sql}
        GROUP BY {sid_expr}
        ORDER BY course_name, section_id
        """,
        tuple(dept_params),
    ).fetchall()

    out: list[dict[str, Any]] = []
    for row in rows:
        if hasattr(row, "keys"):
            d = dict(row)
        else:
            d = {
                "section_id": row[0],
                "course_name": row[1],
                "instructor_id": row[2],
                "instructor_name": row[3],
                "schedule_semester": row[4],
                "department_name": row[5],
            }
        sid = int(d.get("section_id") or 0)
        if not sid:
            continue
        sch_sem = (d.get("schedule_semester") or "").strip()
        if sch_sem and sem and not schedule_semester_matches_current_term(sch_sem, sem):
            continue
        iid = _resolve_schedule_instructor_id(
            int(d.get("instructor_id") or 0),
            str(d.get("instructor_name") or ""),
            name_map,
        )
        cname = (d.get("course_name") or "").strip()
        if not cname:
            continue
        out.append(
            {
                "section_id": sid,
                "course_name": cname,
                "instructor_id": iid,
                "instructor_name": (d.get("instructor_name") or "").strip() or "—",
                "schedule_semester": sch_sem,
                "department_name": (d.get("department_name") or "").strip() or "—",
            }
        )
    return out


def _course_eval_gap_reasons(
    *,
    has_instructor: bool,
    registration_count: int,
    schedule_semester: str,
    term: str,
) -> list[str]:
    reasons = ["لم يُرسَل أي تقييم"]
    if not has_instructor:
        reasons.append("بلا أستاذ معيّن في الجدول")
    if registration_count <= 0:
        reasons.append("بلا تسجيلات طلاب للمقرر")
    if term and not (schedule_semester or "").strip():
        reasons.append("حقل فصل الشعبة فارغ في الجدول")
    return reasons


def build_course_eval_missing_sections_audit(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
) -> dict[str, Any]:
    """
    تدقيق شعب الجدول للفصل التي لم يُرسَل لها أي تقييم مقرر.
    يُستخدم في صفحة النتائج لتبيين الفجوة بين الجدول والاستبيان.
    """
    sem = (semester or "").strip() or term_label_from_conn(conn)
    from backend.services import teaching_groups as tg_svc

    if tg_svc.semester_has_teaching_groups(conn, sem):
        tg_audit = tg_svc.teaching_groups_without_evaluation_audit(
            conn, semester=sem, department_id=department_id
        )
        rows = []
        for r in tg_audit.get("rows") or []:
            rows.append(
                {
                    "section_id": tg_svc.primary_section_id_for_group(
                        conn, int(r.get("teaching_group_id") or 0)
                    ),
                    "teaching_group_id": r.get("teaching_group_id"),
                    "course_name": r.get("course_name"),
                    "instructor_id": r.get("instructor_id"),
                    "instructor_name": r.get("instructor_name"),
                    "department_name": r.get("department_name"),
                    "group_code_label": r.get("group_code_label"),
                    "display_label": r.get("display_label"),
                    "schedule_semester": sem,
                    "course_registration_count": r.get("enrolled_count"),
                    "enrolled_count": r.get("enrolled_count"),
                    "eligible_for_student": r.get("eligible_for_student"),
                    "gap_reasons": r.get("gap_reasons"),
                    "gap_reasons_ar": r.get("gap_reasons_ar"),
                }
            )
        return {
            "semester": sem,
            "department_label": _department_label(conn, department_id),
            "total_schedule_sections": tg_audit.get("total_teaching_groups"),
            "evaluated_sections": tg_audit.get("evaluated_groups"),
            "missing_sections": tg_audit.get("missing_groups"),
            "rows": rows,
            "audit_mode": "teaching_groups",
        }
    schedule_sections = _list_schedule_sections_for_term(conn, sem, department_id=department_id)
    evaluated_ids = _evaluated_section_ids(conn, sem, department_id=department_id)
    evaluated_course_instructor: set[tuple[str, int]] = set()
    if table_exists(conn, "course_evaluations"):
        cur = conn.cursor()
        dept_sql = ""
        params: list[Any] = [sem]
        if department_id is not None:
            pk = schedule_pk_column(conn)
            dept_sql = f"""
                AND EXISTS (
                    SELECT 1 FROM schedule sch
                    WHERE sch.{pk} = e.section_id AND sch.department_id = ?
                )
            """
            params.append(int(department_id))
        for row in cur.execute(
            f"""
            SELECT DISTINCT lower(trim(e.course_name)), e.instructor_id
            FROM course_evaluations e
            WHERE e.semester = ? {dept_sql}
            """,
            tuple(params),
        ).fetchall():
            ckey = (_row_val(row, 0) or "").strip().lower()
            eiid = int(_row_val(row, 1, "instructor_id") or 0)
            if ckey and eiid > 0:
                evaluated_course_instructor.add((ckey, eiid))
    missing_rows: list[dict[str, Any]] = []

    for sec in schedule_sections:
        sid = int(sec["section_id"])
        if sid in evaluated_ids:
            continue
        cname = sec["course_name"]
        reg_count = _course_registration_count(conn, cname)
        iid = int(sec.get("instructor_id") or 0)
        ckey = cname.strip().lower()
        if iid > 0 and (ckey, iid) in evaluated_course_instructor:
            continue
        n_groups = _course_teaching_group_count(conn, cname, sem, instructor_id=iid or None)
        enrolled = section_enrolled_count(conn, cname, sem, section_count=n_groups)
        has_instructor = iid > 0
        gap_reasons = _course_eval_gap_reasons(
            has_instructor=has_instructor,
            registration_count=reg_count,
            schedule_semester=str(sec.get("schedule_semester") or ""),
            term=sem,
        )
        eligible = has_instructor and reg_count > 0
        missing_rows.append(
            {
                **sec,
                "course_registration_count": reg_count,
                "enrolled_count": enrolled,
                "eligible_for_student": eligible,
                "gap_reasons": gap_reasons,
                "gap_reasons_ar": "؛ ".join(gap_reasons),
            }
        )

    return {
        "semester": sem,
        "department_label": _department_label(conn, department_id),
        "total_schedule_sections": len(schedule_sections),
        "evaluated_sections": len(evaluated_ids & {int(s["section_id"]) for s in schedule_sections}),
        "missing_sections": len(missing_rows),
        "rows": missing_rows,
    }


def course_eval_missing_audit_excel_frames(audit: dict[str, Any]) -> list[tuple[str, pd.DataFrame]]:
    """أوراق Excel لتقرير شعب بلا تقييم."""
    rows = [
        {
            "المقرر": r.get("course_name"),
            "الشعبة": r.get("section_id"),
            "الأستاذ": r.get("instructor_name"),
            "القسم": r.get("department_name"),
            "فصل_الجدول": r.get("schedule_semester") or "—",
            "مسجّلون_تقدير": r.get("enrolled_count"),
            "إجمالي_تسجيل_المقرر": r.get("course_registration_count"),
            "يظهر_للطالب": "نعم" if r.get("eligible_for_student") else "لا",
            "أسباب_الفجوة": r.get("gap_reasons_ar"),
        }
        for r in audit.get("rows") or []
    ]
    summary = [
        {"البند": "الفصل", "القيمة": audit.get("semester")},
        {"البند": "النطاق", "القيمة": audit.get("department_label")},
        {"البند": "شعب الجدول للفصل", "القيمة": audit.get("total_schedule_sections")},
        {"البند": "شعب لديها تقييم", "القيمة": audit.get("evaluated_sections")},
        {"البند": "شعب بلا أي تقييم", "القيمة": audit.get("missing_sections")},
        {
            "البند": "ملاحظة",
            "القيمة": (
                "الشعب المدرجة بلا تقييم لا تظهر في جدول النتائج المجمّعة "
                "حتى يُرسَل تقييم واحد على الأقل."
            ),
        },
    ]
    return [
        ("شعب_بلا_تقييم", pd.DataFrame(rows) if rows else pd.DataFrame(columns=["المقرر"])),
        ("ملخص", pd.DataFrame(summary)),
    ]


def export_course_eval_missing_sections_xlsx(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
):
    audit = build_course_eval_missing_sections_audit(
        conn, semester=semester, department_id=department_id
    )
    sem_slug = (audit.get("semester") or "report").replace(" ", "_")[:40]
    return excel_response_from_frames(
        course_eval_missing_audit_excel_frames(audit),
        filename_prefix=f"course_eval_missing_{sem_slug}",
    )


def _aggregate_course_eval_questions(
    conn,
    *,
    semester: str,
    where_sql: str,
    params: list[Any],
) -> list[dict]:
    cur = conn.cursor()
    rating_expr = _eval_rating_subquery()
    use_dynamic = table_exists(conn, "evaluation_survey_answers")
    from backend.services.evaluation_survey import list_survey_questions

    questions_out: list[dict] = []
    for q in list_survey_questions(conn, active_only=True):
        qid = int(q["id"])
        if use_dynamic:
            avg_row = cur.execute(
                f"""
                SELECT AVG(a.rating * 1.0)
                FROM evaluation_survey_answers a
                JOIN course_evaluations e ON e.id = a.evaluation_id
                WHERE e.semester = ? AND a.question_id = ? {where_sql}
                """,
                tuple([semester, qid] + params),
            ).fetchone()
        else:
            avg_row = cur.execute(
                f"""
                SELECT AVG({rating_expr})
                FROM course_evaluations e
                WHERE e.semester = ? {where_sql}
                """,
                tuple([semester] + params),
            ).fetchone()
        avg5 = float((_row_val(avg_row, 0) or 0) or 0)
        pct = round((avg5 / 5.0) * 100.0, 1) if avg5 else None
        questions_out.append(
            {
                "question_id": qid,
                "label_ar": q.get("label_ar"),
                "avg_rating": round(avg5, 2) if avg5 else None,
                "score_percent": pct,
            }
        )
    return _enrich_questions(questions_out)


def _overall_course_eval_score(
    conn,
    *,
    semester: str,
    where_sql: str,
    params: list[Any],
) -> float | None:
    cur = conn.cursor()
    rating_expr = _eval_rating_subquery()
    row = cur.execute(
        f"""
        SELECT AVG({rating_expr})
        FROM course_evaluations e
        WHERE e.semester = ? {where_sql}
        """,
        tuple([semester] + params),
    ).fetchone()
    avg5 = float((_row_val(row, 0) or 0) or 0)
    return round((avg5 / 5.0) * 100.0, 1) if avg5 else None


def _cache_get(cache: dict[str, Any] | None, key: Any, factory):
    if cache is None:
        return factory()
    if key not in cache:
        cache[key] = factory()
    return cache[key]


def _section_groups_cached(
    conn,
    sem: str,
    department_id: int | None,
    *,
    section_groups: list[dict[str, Any]] | None = None,
    eval_cache: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    if section_groups is not None:
        return section_groups
    key = ("section_groups", sem, department_id)
    return _cache_get(
        eval_cache,
        key,
        lambda: _fetch_course_eval_section_groups(conn, sem, department_id),
    )


def _course_eval_columns_cached(
    conn,
    *,
    eval_cache: dict[str, Any] | None = None,
) -> set[str]:
    return _cache_get(
        eval_cache,
        ("course_eval_columns",),
        lambda: {c.lower() for c in fetch_table_columns(conn, "course_evaluations")},
    )


def _course_eval_finalize_meta_cached(
    conn,
    *,
    semester: str,
    department_id: int | None = None,
    eval_cache: dict[str, Any] | None = None,
) -> dict[str, Any]:
    key = ("finalize_meta", "student_course", semester, department_id)

    def _build() -> dict[str, Any]:
        return {
            "accreditation_links": accreditation_links_for(
                "student_course", conn, semester=semester, department_id=department_id
            ),
            "primary_accreditation": _primary_accreditation_label(
                conn, "student_course", semester=semester, department_id=department_id
            ),
            "course_eval_policy_ar": format_course_eval_aggregation_policy(conn),
        }

    return _cache_get(eval_cache, key, _build)


def _finalize_course_eval_unit_report(
    report: dict[str, Any],
    conn,
    *,
    semester: str,
    department_id: int | None = None,
    title_suffix: str = "",
    eval_cache: dict[str, Any] | None = None,
    summary_only: bool = False,
) -> dict[str, Any]:
    questions = report.get("questions") or []
    weakest, strongest = _weakest_strongest(questions)
    score = report.get("overall_score_percent")
    aggregated = bool(report.get("aggregated"))
    compliance = classify_compliance_status(score if aggregated else None)
    base_title = "تقييم المقرر والأستاذ (طالب)"
    title = f"{base_title}{title_suffix}".strip()
    enrolled = int(report.get("enrolled_count") or 0)
    meta = _course_eval_finalize_meta_cached(
        conn, semester=semester, department_id=department_id, eval_cache=eval_cache
    )
    if summary_only:
        report.update(
            {
                "template_code": "student_course",
                "title_ar": title,
                "min_aggregate": report.get("min_aggregate")
                or course_eval_min_required(enrolled, conn=conn),
                "course_eval_policy_ar": meta["course_eval_policy_ar"],
                "questions": [],
                "weakest_item": "—",
                "strongest_item": "—",
                "compliance_status": compliance,
                "compliance_status_ar": COMPLIANCE_STATUS_LABELS_AR.get(compliance, compliance),
                "accreditation_links": [],
                "primary_accreditation": "—",
                "respondent_label": RESPONDENT_ROLE_LABELS.get("student", "الطالب"),
                "recommendations": [],
            }
        )
        return report
    report.update(
        {
            "template_code": "student_course",
            "title_ar": title,
            "min_aggregate": report.get("min_aggregate")
            or course_eval_min_required(enrolled, conn=conn),
            "course_eval_policy_ar": meta["course_eval_policy_ar"],
            "questions": questions,
            "weakest_item": weakest,
            "strongest_item": strongest,
            "compliance_status": compliance,
            "compliance_status_ar": COMPLIANCE_STATUS_LABELS_AR.get(compliance, compliance),
            "accreditation_links": meta["accreditation_links"],
            "primary_accreditation": meta["primary_accreditation"],
            "respondent_label": RESPONDENT_ROLE_LABELS.get("student", "الطالب"),
            "recommendations": generate_recommendations(questions, title),
        }
    )
    return report


def build_course_eval_section_report(
    conn,
    section_id: int,
    *,
    semester: str | None = None,
    department_id: int | None = None,
    course_name: str | None = None,
    instructor_id: int | None = None,
    teaching_group_id: int | None = None,
    group_meta: dict[str, Any] | None = None,
    section_groups: list[dict[str, Any]] | None = None,
    eval_cache: dict[str, Any] | None = None,
    summary_only: bool = False,
) -> dict[str, Any] | None:
    """تقرير تجميعي لتقييم شعبة أو مجموعة تدريس."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    sid = int(section_id)
    tgid = int(teaching_group_id or 0)
    group = group_meta
    if not group:
        groups = _section_groups_cached(
            conn, sem, department_id, section_groups=section_groups, eval_cache=eval_cache
        )
        if tgid > 0:
            group = next((g for g in groups if int(g.get("teaching_group_id") or 0) == tgid), None)
        if not group:
            group = next((g for g in groups if int(g["section_id"]) == sid), None)
    if not group and course_name and instructor_id:
        group = {
            "section_id": sid,
            "course_name": course_name.strip(),
            "instructor_id": int(instructor_id),
            "instructor_name": "—",
            "department_name": "—",
            "response_count": 0,
        }
    if not group:
        return None

    count = int(group["response_count"])
    iid = int(group["instructor_id"])
    tgid = int(group.get("teaching_group_id") or 0)
    cname = group["course_name"]
    reg_total = _cache_get(
        eval_cache,
        ("reg_count", cname),
        lambda: _course_registration_count(conn, cname),
    )
    n_groups = _cache_get(
        eval_cache,
        ("tg_count", cname, sem, iid),
        lambda: _course_teaching_group_count(conn, cname, sem, instructor_id=iid),
    )
    enrolled = section_enrolled_count(
        conn,
        cname,
        sem,
        section_count=n_groups,
        teaching_group_id=tgid or None,
    )
    min_req = course_eval_min_required(enrolled, conn=conn)
    aggregated = course_eval_is_aggregated(count, enrolled, conn=conn)
    ce_cols = _course_eval_columns_cached(conn, eval_cache=eval_cache)
    if tgid > 0 and "teaching_group_id" in ce_cols:
        where_sql = " AND e.teaching_group_id = ? AND e.course_name = ? AND e.instructor_id = ?"
        params = [tgid, group["course_name"], iid]
    else:
        where_sql = " AND e.section_id = ? AND e.course_name = ? AND e.instructor_id = ?"
        params = [sid, group["course_name"], iid]

    questions: list[dict] = []
    overall = None
    if aggregated:
        overall = _overall_course_eval_score(conn, semester=sem, where_sql=where_sql, params=params)
        if not summary_only:
            questions = _aggregate_course_eval_questions(
                conn, semester=sem, where_sql=where_sql, params=params
            )

    report = {
        "section_id": sid,
        "teaching_group_id": tgid or None,
        "group_code_label": group.get("group_code_label"),
        "course_name": group["course_name"],
        "instructor_id": group["instructor_id"],
        "instructor_name": group["instructor_name"],
        "department_name": group["department_name"],
        "semester": sem,
        "enrolled_count": enrolled,
        "course_registration_count": reg_total,
        "teaching_group_count": n_groups,
        "min_aggregate": min_req,
        "response_rate_percent": round((count / enrolled) * 100.0, 1) if enrolled > 0 else None,
        "response_count": count,
        "aggregated": aggregated,
        "overall_score_percent": overall,
        "questions": questions,
        "group_type": "section",
        "section_ids": [sid],
        "section_count": 1,
    }
    return _finalize_course_eval_unit_report(
        report,
        conn,
        semester=sem,
        department_id=department_id,
        title_suffix=f" — {group['course_name']} (شعبة {sid})",
        eval_cache=eval_cache,
        summary_only=summary_only,
    )


def build_course_eval_by_course_report(
    conn,
    course_name: str,
    instructor_id: int,
    *,
    semester: str | None = None,
    department_id: int | None = None,
    section_groups: list[dict[str, Any]] | None = None,
    eval_cache: dict[str, Any] | None = None,
) -> dict[str, Any] | None:
    """تجميع تقييمات المقرر لنفس الأستاذ عبر كل شعبِه."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    cname = (course_name or "").strip()
    iid = int(instructor_id)
    if not cname or not iid:
        return None

    all_groups = _section_groups_cached(
        conn, sem, department_id, section_groups=section_groups, eval_cache=eval_cache
    )
    groups = [
        g
        for g in all_groups
        if int(g.get("instructor_id") or 0) == iid
        and (g.get("course_name") or "").strip().lower() == cname.lower()
    ]
    if not groups:
        return None

    count = sum(int(g["response_count"]) for g in groups)
    enrolled = _course_registration_count(conn, cname)
    min_req = course_eval_min_required(enrolled, conn=conn)
    aggregated = course_eval_is_aggregated(count, enrolled, conn=conn)
    dept_sql, dept_params = _course_eval_dept_filter(conn, department_id)
    where_sql = (
        " AND lower(trim(e.course_name)) = lower(trim(?)) AND e.instructor_id = ?"
        + dept_sql
    )
    params = [cname, iid] + dept_params

    questions: list[dict] = []
    overall = None
    if aggregated:
        overall = _overall_course_eval_score(conn, semester=sem, where_sql=where_sql, params=params)
        questions = _aggregate_course_eval_questions(
            conn, semester=sem, where_sql=where_sql, params=params
        )

    instructor_name = (groups[0].get("instructor_name") or "—").strip()
    department_names = sorted({(g.get("department_name") or "—") for g in groups})
    section_ids = [int(g["section_id"]) for g in groups]

    report = {
        "section_id": section_ids[0] if len(section_ids) == 1 else None,
        "course_name": cname,
        "instructor_id": iid,
        "instructor_name": instructor_name,
        "department_name": "، ".join(department_names),
        "semester": sem,
        "enrolled_count": enrolled,
        "min_aggregate": min_req,
        "response_rate_percent": round((count / enrolled) * 100.0, 1) if enrolled > 0 else None,
        "response_count": count,
        "aggregated": aggregated,
        "overall_score_percent": overall,
        "questions": questions,
        "group_type": "course_instructor",
        "section_ids": section_ids,
        "section_count": len(section_ids),
    }
    suffix = f" — {cname} / {instructor_name}"
    if len(section_ids) > 1:
        suffix += f" ({len(section_ids)} شعب)"
    return _finalize_course_eval_unit_report(
        report,
        conn,
        semester=sem,
        department_id=department_id,
        title_suffix=suffix,
        eval_cache=eval_cache,
    )


def build_course_eval_sections_summary(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
    section_groups: list[dict[str, Any]] | None = None,
    eval_cache: dict[str, Any] | None = None,
    summary_only: bool = False,
) -> list[dict[str, Any]]:
    """ملخص تقييم لكل شعبة."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    groups = _section_groups_cached(
        conn, sem, department_id, section_groups=section_groups, eval_cache=eval_cache
    )
    reports: list[dict[str, Any]] = []
    for g in groups:
        rep = build_course_eval_section_report(
            conn,
            int(g["section_id"]),
            semester=sem,
            department_id=department_id,
            group_meta=g,
            section_groups=groups,
            eval_cache=eval_cache,
            summary_only=summary_only,
        )
        if rep:
            reports.append(rep)
    return reports


def list_course_eval_course_instructor_groups(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
    section_reports: list[dict[str, Any]] | None = None,
    section_groups: list[dict[str, Any]] | None = None,
    eval_cache: dict[str, Any] | None = None,
) -> list[dict[str, Any]]:
    """مجموعات مقرر+أستاذ (شعب متعددة مجمّعة)."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    groups = _section_groups_cached(
        conn, sem, department_id, section_groups=section_groups, eval_cache=eval_cache
    )
    if section_reports is None:
        section_reports = build_course_eval_sections_summary(
            conn,
            semester=sem,
            department_id=department_id,
            section_groups=groups,
            eval_cache=eval_cache,
        )
    by_key: dict[tuple[str, int], dict[str, Any]] = {}
    for r in section_reports:
        key = ((r.get("course_name") or "").strip().lower(), int(r.get("instructor_id") or 0))
        if not key[0] or not key[1]:
            continue
        bucket = by_key.get(key)
        if not bucket:
            by_key[key] = {
                "course_name": r.get("course_name"),
                "instructor_id": key[1],
                "instructor_name": r.get("instructor_name"),
                "department_name": r.get("department_name"),
                "section_ids": [int(r["section_id"])],
                "response_count": int(r.get("response_count") or 0),
                "section_count": 1,
            }
        else:
            bucket["section_ids"].append(int(r["section_id"]))
            bucket["response_count"] += int(r.get("response_count") or 0)
            bucket["section_count"] = len(bucket["section_ids"])
            depts = {bucket.get("department_name"), r.get("department_name")}
            bucket["department_name"] = "، ".join(sorted(d for d in depts if d and d != "—"))

    out: list[dict[str, Any]] = []
    for (_c, _i), meta in sorted(by_key.items(), key=lambda x: (x[1]["course_name"] or "")):
        if int(meta.get("section_count") or 0) <= 1:
            continue
        full = build_course_eval_by_course_report(
            conn,
            meta["course_name"],
            int(meta["instructor_id"]),
            semester=sem,
            department_id=department_id,
            section_groups=groups,
            eval_cache=eval_cache,
        )
        if full:
            out.append(full)
    return out


def build_course_eval_results_bundle(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
    summary_only: bool = False,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """ملخص الشعب + تجميع مقرر/أستاذ في تمرير واحد مع cache مشترك."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    eval_cache: dict[str, Any] = {}
    groups = _fetch_course_eval_section_groups(conn, sem, department_id)
    eval_cache[("section_groups", sem, department_id)] = groups
    sections = build_course_eval_sections_summary(
        conn,
        semester=sem,
        department_id=department_id,
        section_groups=groups,
        eval_cache=eval_cache,
        summary_only=summary_only,
    )
    by_course = list_course_eval_course_instructor_groups(
        conn,
        semester=sem,
        department_id=department_id,
        section_groups=groups,
        section_reports=sections,
        eval_cache=eval_cache,
    )
    return sections, by_course


def build_course_eval_report(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
) -> dict[str, Any]:
    """تقرير تجميعي لتقييم المقرر (مسار student_course)."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    cur = conn.cursor()
    dept_sql = ""
    params: list[Any] = [sem]
    if department_id is not None:
        pk = schedule_pk_column(conn)
        dept_sql = f"""
            AND EXISTS (
                SELECT 1 FROM schedule sch
                WHERE sch.{pk} = e.section_id AND sch.department_id = ?
            )
        """
        params.append(int(department_id))

    count_row = cur.execute(
        f"SELECT COUNT(*) FROM course_evaluations e WHERE e.semester = ? {dept_sql}",
        tuple(params),
    ).fetchone()
    count = int((count_row[0] if count_row else 0) or 0)
    enrolled = _college_course_eval_enrolled(conn, sem, department_id)
    min_n = course_eval_min_required(enrolled, conn=conn)
    aggregated = course_eval_is_aggregated(count, enrolled, conn=conn)
    overall = _avg_eval_score(conn, cur, sem, department_id) if aggregated else None

    questions: list[dict] = []
    if aggregated and table_exists(conn, "evaluation_survey_answers"):
        from backend.services.evaluation_survey import list_survey_questions

        for q in list_survey_questions(conn, active_only=True):
            qid = int(q["id"])
            avg_row = cur.execute(
                f"""
                SELECT AVG(a.rating * 1.0)
                FROM evaluation_survey_answers a
                JOIN course_evaluations e ON e.id = a.evaluation_id
                WHERE e.semester = ? AND a.question_id = ? {dept_sql}
                """,
                tuple([sem, qid] + (params[1:] if len(params) > 1 else [])),
            ).fetchone()
            avg5 = float((avg_row[0] if avg_row else 0) or 0)
            pct = round((avg5 / 5.0) * 100.0, 1) if avg5 else None
            questions.append(
                {
                    "question_id": qid,
                    "label_ar": q.get("label_ar"),
                    "avg_rating": round(avg5, 2) if avg5 else None,
                    "score_percent": pct,
                }
            )
        questions = _enrich_questions(questions)

    compliance = classify_compliance_status(overall if aggregated else None)
    return {
        "template_code": "student_course",
        "title_ar": "تقييم المقرر والأستاذ (طالب)",
        "semester": sem,
        "response_count": count,
        "enrolled_count": enrolled,
        "min_aggregate": min_n,
        "response_rate_percent": round((count / enrolled) * 100.0, 1) if enrolled > 0 else None,
        "aggregated": aggregated,
        "overall_score_percent": round(overall, 1) if overall is not None else None,
        "questions": questions,
        "compliance_status": compliance,
        "compliance_status_ar": COMPLIANCE_STATUS_LABELS_AR.get(compliance, compliance),
        "accreditation_links": accreditation_links_for(
            "student_course", conn, semester=sem, department_id=department_id
        ),
        "primary_accreditation": _primary_accreditation_label(
            conn, "student_course", semester=sem, department_id=department_id
        ),
        "respondent_label": RESPONDENT_ROLE_LABELS.get("student", "الطالب"),
        "course_eval_policy_ar": format_course_eval_aggregation_policy(conn),
    }


def build_combined_survey_report(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
    include_course_eval: bool = True,
) -> dict[str, Any]:
    """تقرير موحّد لكل استبيانات المنصة (+ تقييم المقرر اختيارياً)."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    reports: list[dict[str, Any]] = []
    for t in list_templates(conn):
        if int(t.get("legacy_course_eval") or 0):
            continue
        reports.append(
            build_survey_report(conn, t["code"], semester=sem, department_id=department_id)
        )
    course_eval = None
    course_eval_sections: list[dict[str, Any]] = []
    course_eval_by_course: list[dict[str, Any]] = []
    if include_course_eval:
        course_eval = build_course_eval_report(conn, semester=sem, department_id=department_id)
        course_eval_sections, course_eval_by_course = build_course_eval_results_bundle(
            conn, semester=sem, department_id=department_id
        )

    aggregated_count = sum(1 for r in reports if r.get("aggregated"))
    if course_eval and course_eval.get("aggregated"):
        aggregated_count += 1

    scored = [
        r
        for r in reports
        if r.get("aggregated") and r.get("overall_score_percent") is not None
    ]
    if course_eval and course_eval.get("aggregated") and course_eval.get("overall_score_percent") is not None:
        scored.append(course_eval)

    top3 = sorted(scored, key=lambda x: float(x["overall_score_percent"]), reverse=True)[:3]
    bottom3 = sorted(scored, key=lambda x: float(x["overall_score_percent"]))[:3]

    return {
        "semester": sem,
        "department_id": department_id,
        "department_label": _department_label(conn, department_id),
        "generated_at": datetime.datetime.utcnow().isoformat(timespec="seconds"),
        "reports": reports,
        "course_eval": course_eval,
        "course_eval_sections": course_eval_sections,
        "course_eval_by_course": course_eval_by_course,
        "aggregated_survey_count": aggregated_count,
        "total_survey_count": len(reports) + (1 if include_course_eval else 0),
        "top_surveys": top3,
        "bottom_surveys": bottom3,
        "course_eval_policy_ar": format_course_eval_aggregation_policy(conn),
    }


def _accreditation_map_rows(reports: list[dict], course_eval: dict | None) -> list[dict]:
    rows: list[dict] = []
    all_reports = list(reports)
    if course_eval:
        all_reports.append(course_eval)
    for r in all_reports:
        code = r.get("template_code") or ""
        for link in r.get("accreditation_links") or []:
            rows.append(
                {
                    "الاستبيان": r.get("title_ar"),
                    "رمز_الاستبيان": code,
                    "المعيار": link.get("standard_code"),
                    "المؤشر": link.get("indicator_code"),
                    "عنوان_المؤشر": link.get("indicator_title_ar"),
                    "نوع_الربط": LINK_TYPE_LABELS_AR.get(
                        link.get("link_type", ""), link.get("link_type")
                    ),
                    "كيفية_الاستفادة": link.get("usage_ar"),
                    "نتيجة_الاستبيان_%": r.get("overall_score_percent"),
                    "حالة_الامتثال": r.get("compliance_status_ar"),
                }
            )
    return rows


def _executive_summary_rows(reports: list[dict], course_eval: dict | None) -> list[dict]:
    rows: list[dict] = []
    for r in reports:
        rows.append(_summary_row(r))
    if course_eval:
        rows.append(_summary_row(course_eval))
    return rows


def _summary_row(r: dict) -> dict:
    return {
        "الاستبيان": r.get("title_ar"),
        "الرمز": r.get("template_code"),
        "الفئة": r.get("respondent_label"),
        "عدد_الإجابات": r.get("response_count"),
        "الحد_الأدنى": r.get("min_aggregate"),
        "حالة_التجميع": "مكتمل" if r.get("aggregated") else "ناقص",
        "النتيجة_%": r.get("overall_score_percent"),
        "أضعف_بند": r.get("weakest_item", "—"),
        "أقوى_بند": r.get("strongest_item", "—"),
        "معيار_الاعتماد": r.get("primary_accreditation"),
        "حالة_الامتثال": r.get("compliance_status_ar"),
    }


def _comparative_analysis_rows(reports: list[dict], course_eval: dict | None) -> list[dict]:
    rows: list[dict] = []
    by_role: dict[str, list[float]] = {}
    all_reports = list(reports)
    if course_eval:
        all_reports.append(course_eval)

    for r in all_reports:
        if not r.get("aggregated") or r.get("overall_score_percent") is None:
            continue
        role = r.get("respondent_label") or "—"
        by_role.setdefault(role, []).append(float(r["overall_score_percent"]))

    for role, scores in sorted(by_role.items()):
        rows.append(
            {
                "نوع_التحليل": "متوسط حسب الفئة",
                "الفئة": role,
                "القيمة": round(sum(scores) / len(scores), 1) if scores else None,
                "التفاصيل": f"{len(scores)} استبيان(ات)",
            }
        )

    ranked = [
        r
        for r in all_reports
        if r.get("aggregated") and r.get("overall_score_percent") is not None
    ]
    ranked.sort(key=lambda x: float(x["overall_score_percent"]), reverse=True)
    for i, r in enumerate(ranked, 1):
        rows.append(
            {
                "نوع_التحليل": "ترتيب الاستبيانات",
                "الفئة": f"#{i}",
                "القيمة": r.get("overall_score_percent"),
                "التفاصيل": f"{r.get('title_ar')} ({r.get('template_code')})",
            }
        )

    for r in all_reports:
        if not r.get("aggregated"):
            continue
        for q in r.get("questions") or []:
            pct = q.get("score_percent")
            if pct is not None and float(pct) < 60:
                rows.append(
                    {
                        "نوع_التحليل": "بند تحت 60%",
                        "الفئة": r.get("title_ar"),
                        "القيمة": pct,
                        "التفاصيل": (q.get("label_ar") or "")[:100],
                    }
                )
            elif pct is not None and float(pct) >= 80:
                rows.append(
                    {
                        "نوع_التحليل": "بند فوق 80%",
                        "الفئة": r.get("title_ar"),
                        "القيمة": pct,
                        "التفاصيل": (q.get("label_ar") or "")[:100],
                    }
                )
    return rows


def _question_rows(report: dict) -> list[dict]:
    rows: list[dict] = []
    for i, q in enumerate(report.get("questions") or [], 1):
        rows.append(
            {
                "ترتيب": i,
                "البند": q.get("label_ar"),
                "متوسط_1_5": q.get("avg_rating"),
                "النسبة_%": q.get("score_percent"),
                "التصنيف": q.get("classification_ar"),
                "توصية": q.get("recommendation_ar"),
            }
        )
    return rows


def _metadata_rows(combined: dict) -> list[dict]:
    return [
        {"البند": "الفصل الدراسي", "القيمة": combined.get("semester")},
        {"البند": "النطاق", "القيمة": combined.get("department_label")},
        {"البند": "تاريخ التصدير (UTC)", "القيمة": combined.get("generated_at")},
        {"البند": "عدد الاستبيانات", "القيمة": combined.get("total_survey_count")},
        {"البند": "استبيانات مجمّعة", "القيمة": combined.get("aggregated_survey_count")},
        {
            "البند": "سياسة الخصوصية",
            "القيمة": "لا تُصدَّر إجابات فردية — التجميع بعد الحد الأدنى فقط.",
        },
        {"البند": "عتبة المتحقق", "القيمة": "≥ 70%"},
        {"البند": "عتبة الجزئي", "القيمة": "50% – 69%"},
        {"البند": "مقياس التقييم", "القيمة": "Likert 1–5 → نسبة = (متوسط/5)×100"},
        {
            "البند": "حد تجميع تقييم المقرر",
            "القيمة": combined.get("course_eval_policy_ar")
            or format_course_eval_aggregation_policy(),
        },
    ]


def _sheet_name_for_code(code: str, title_ar: str = "") -> str:
    short = {
        "student_services": "خدمات_الطالب",
        "student_facilities": "مرافق_الطالب",
        "faculty_hod": "رئيس_القسم",
        "faculty_dean": "قيادة_الكلية_وسياساتها",
        "faculty_educational_process": "العملية_التعليمية",
        "supervisor_advising": "مشرف_ارشاد",
        "supervisor_coordination": "مشرف_تنسيق",
        "staff_workplace": "موظف_بيئة",
        "staff_student_services": "موظف_خدمة",
        "student_course": "تقييم_المقرر",
    }
    return short.get(code, (title_ar or code)[:28])


def _course_eval_sections_export_stats(sections: list[dict]) -> dict[str, Any]:
    """إحصاءات ملخصة لتصدير الشعب (بدون ربط اعتماد)."""
    aggregated = [s for s in sections if s.get("aggregated")]
    pending = [s for s in sections if not s.get("aggregated")]
    scores = [
        float(s["overall_score_percent"])
        for s in aggregated
        if s.get("overall_score_percent") is not None
    ]
    avg_score = round(sum(scores) / len(scores), 1) if scores else None
    return {
        "section_count": len(sections),
        "aggregated_count": len(aggregated),
        "pending_count": len(pending),
        "avg_score_percent": avg_score,
    }


def _aggregate_department_question_scores(sections: list[dict]) -> list[dict[str, Any]]:
    """متوسط بنود الاستبيان عبر الشعب المجمّعة في القسم."""
    by_label: dict[str, list[float]] = {}
    for s in sections:
        if not s.get("aggregated"):
            continue
        for q in s.get("questions") or []:
            pct = q.get("score_percent")
            if pct is None:
                continue
            label = (q.get("label_ar") or "").strip()
            if not label:
                continue
            by_label.setdefault(label, []).append(float(pct))
    out: list[dict[str, Any]] = []
    for label, scores in by_label.items():
        avg = round(sum(scores) / len(scores), 1)
        cls = classify_item_score(avg)
        out.append(
            {
                "label_ar": label,
                "score_percent": avg,
                "classification": cls,
                "classification_ar": SCORE_CLASS_LABELS_AR.get(cls, cls),
                "section_count": len(scores),
            }
        )
    out.sort(key=lambda x: float(x["score_percent"]), reverse=True)
    return out


def _section_label_short(s: dict) -> str:
    course = (s.get("course_name") or "—").strip()
    sec = s.get("section_id")
    return f"{course} (ش.{sec})" if sec is not None else course


def build_course_eval_sections_analysis(
    *,
    sections: list[dict],
    stats: dict[str, Any],
    missing_audit: dict[str, Any] | None,
    department_label: str,
    semester: str,
    course_eval_policy_ar: str,
) -> dict[str, Any]:
    """تحليل واستنتاجات آلية لتقرير الشعب (شاهد)."""
    aggregated = [s for s in sections if s.get("aggregated") and s.get("overall_score_percent") is not None]
    pending = [s for s in sections if not s.get("aggregated")]
    missing_rows = (missing_audit or {}).get("rows") or []
    avg = stats.get("avg_score_percent")
    interpretation = interpret_overall_score_ar(avg, bool(aggregated))

    dist_labels = dict(SECTION_SCORE_LABELS_AR)
    buckets: dict[str, list[dict]] = {key: [] for key in SECTION_SCORE_BUCKET_ORDER}
    for s in sections:
        if not s.get("aggregated") or s.get("overall_score_percent") is None:
            buckets["pending"].append(s)
            continue
        buckets[classify_section_score(float(s["overall_score_percent"]))].append(s)

    distribution_rows = [
        {
            "التصنيف": dist_labels[key],
            "عدد_الشعب": len(bucket),
            "أمثلة": "؛ ".join(_section_label_short(s) for s in bucket[:3]) or "—",
        }
        for key in SECTION_SCORE_BUCKET_ORDER
        for bucket in [buckets[key]]
        if bucket
    ]

    sorted_agg = sorted(aggregated, key=lambda x: float(x["overall_score_percent"]), reverse=True)
    top_sections = [
        {
            "المقرر": s.get("course_name"),
            "الشعبة": s.get("section_id"),
            "الأستاذ": s.get("instructor_name"),
            "النتيجة_%": s.get("overall_score_percent"),
        }
        for s in sorted_agg[:3]
    ]
    bottom_sections = [
        {
            "المقرر": s.get("course_name"),
            "الشعبة": s.get("section_id"),
            "الأستاذ": s.get("instructor_name"),
            "النتيجة_%": s.get("overall_score_percent"),
        }
        for s in sorted(sorted_agg, key=lambda x: float(x["overall_score_percent"]))[:3]
    ]

    dept_questions = _aggregate_department_question_scores(sections)
    strongest_items = dept_questions[:3]
    weakest_items = list(reversed(dept_questions[-3:])) if dept_questions else []

    follow_up: list[dict[str, Any]] = []
    for s in pending:
        reason = "لم يكتمل عدد المقيّمين"
        if int(s.get("response_count") or 0) == 0:
            reason = "لا يوجد أي تقييم"
        follow_up.append(
            {
                "المقرر": s.get("course_name"),
                "الشعبة": s.get("section_id"),
                "الأستاذ": s.get("instructor_name"),
                "السبب": reason,
            }
        )
    for s in aggregated:
        if float(s["overall_score_percent"]) < 65:
            follow_up.append(
                {
                    "المقرر": s.get("course_name"),
                    "الشعبة": s.get("section_id"),
                    "الأستاذ": s.get("instructor_name"),
                    "السبب": f"نتيجة منخفضة ({s['overall_score_percent']}%)",
                }
            )

    total = int(stats.get("section_count") or 0)
    agg_n = int(stats.get("aggregated_count") or 0)
    pend_n = int(stats.get("pending_count") or 0)
    cov_pct = round((agg_n / total) * 100, 1) if total else 0
    avg_txt = f"{avg}%" if avg is not None else "—"

    narrative: list[str] = [
        (
            f"يغطي هذا التقرير {total} شعبة في قسم «{department_label}» للفصل «{semester}». "
            f"اكتمل التجميع في {agg_n} شعبة ({cov_pct}%) وفق سياسة {course_eval_policy_ar}."
        ),
        (
            f"بلغ متوسط نتائج الشعب المجمّعة {avg_txt}. {interpretation}"
            if aggregated
            else "لا تتوفر نتيجة مجمّعة كافية لإصدار حكم عام على مستوى القسم."
        ),
    ]
    if top_sections:
        best = sorted_agg[0]
        narrative.append(
            f"أعلى نتيجة: «{best.get('course_name')}» شعبة {best.get('section_id')} "
            f"({best.get('overall_score_percent')}%)."
        )
    if len(sorted_agg) > 1:
        worst = sorted(sorted_agg, key=lambda x: float(x["overall_score_percent"]))[0]
        narrative.append(
            f"أدنى نتيجة مجمّعة: «{worst.get('course_name')}» شعبة {worst.get('section_id')} "
            f"({worst.get('overall_score_percent')}%)."
        )
    if strongest_items:
        labels = "، ".join(f"«{q['label_ar']}» ({q['score_percent']}%)" for q in strongest_items)
        narrative.append(f"أبرز نقاط القوة على مستوى القسم: {labels}.")
    if weakest_items:
        labels = "، ".join(f"«{q['label_ar']}» ({q['score_percent']}%)" for q in weakest_items)
        narrative.append(f"أبرز محاور التحسين: {labels}.")
    if missing_rows:
        narrative.append(
            f"وُجدت {len(missing_rows)} شعبة في الجدول الدراسي بلا أي تقييم مقرر — يستحق الأمر متابعة إدارية."
        )

    conclusions: list[str] = []
    if agg_n >= max(1, total // 2):
        conclusions.append("تغطية التقييم على مستوى الشعب مقبولة بشكل عام لهذا الفصل.")
    elif total:
        conclusions.append("تغطية التقييم دون المستوى المطلوب — يُنصح بتعزيز حث الطلبة على التقييم.")
    if avg is not None and avg >= 75:
        conclusions.append("مستوى الرضا العام عن المقررات والتدريس يُعد جيداً ضمن الشعب المجمّعة.")
    elif avg is not None:
        conclusions.append("النتائج العامة تستدعي خطة تحسين فصلية في جودة التدريس والتقييم.")
    if buckets["needs_improvement"]:
        conclusions.append("توجد شعب بنتائج دون 65% تستدعي متابعة فردية مع الأساتذة وخطة تحسين.")
    elif buckets["good"]:
        conclusions.append("بعض الشعب في نطاق «جيد» (65–75%) — يُنصح بمتابعة دورية دون تأخير.")
    if pend_n:
        conclusions.append(f"وجود {pend_n} شعبة بلا تجميع كافٍ يحد من دقة الصورة الكاملة للقسم.")
    if missing_rows:
        conclusions.append("شعب بلا تقييم تُثير تساؤلاً عن اكتمال آلية التقييم — يُراجع مع الشؤون الأكاديمية.")
    if not conclusions:
        conclusions.append("لا تتوفر بيانات كافية لاستنتاجات نهائية — يُعاد التقييم بعد اكتمال التغطية.")

    recommendations: list[dict[str, str]] = []
    for q in weakest_items[:2]:
        recommendations.append(
            {
                "التوصية": f"عقد ورشة أو اجتماع لجنة الجودة حول بند «{q['label_ar']}» ({q['score_percent']}%).",
                "المسؤول": "رئيس القسم / لجنة الجودة",
                "الإطار": "خلال الفصل الحالي",
            }
        )
    for fu in follow_up[:3]:
        recommendations.append(
            {
                "التوصية": (
                    f"متابعة شعبة {fu.get('شعبة')} — {fu.get('المقرر')} "
                    f"({fu.get('السبب')})."
                ),
                "المسؤول": "رئيس القسم",
                "الإطار": "خلال أسبوعين",
            }
        )
    if top_sections and not recommendations:
        recommendations.append(
            {
                "التوصية": "توثيق ممارسات الشعب ذات النتائج العالية كمرجع للتدريس في القسم.",
                "المسؤول": "لجنة الجودة",
                "الإطار": "الفصل القادم",
            }
        )
    if not recommendations:
        recommendations.append(
            {
                "التوصية": "مراجعة نتائج التقييم في اجتماع دوري لضمان الجودة.",
                "المسؤول": "رئيس القسم",
                "الإطار": "نهاية الفصل",
            }
        )

    return {
        "interpretation_ar": interpretation,
        "narrative_paragraphs": narrative,
        "conclusions": conclusions,
        "recommendations": recommendations,
        "distribution_rows": distribution_rows,
        "top_sections": top_sections,
        "bottom_sections": bottom_sections,
        "strongest_items": strongest_items,
        "weakest_items": weakest_items,
        "follow_up_sections": follow_up,
    }


def _analysis_excel_frames(analysis: dict[str, Any]) -> list[tuple[str, pd.DataFrame]]:
    """أوراق التحليل والاستنتاجات."""
    frames: list[tuple[str, pd.DataFrame]] = []
    narrative_rows = [{"الفقرة": p} for p in analysis.get("narrative_paragraphs") or []]
    frames.append(("التحليل", pd.DataFrame(narrative_rows or [{"الفقرة": "—"}])))
    if analysis.get("distribution_rows"):
        frames.append(("توزيع_النتائج", pd.DataFrame(analysis["distribution_rows"])))
    if analysis.get("top_sections"):
        frames.append(("أعلى_الشعب", pd.DataFrame(analysis["top_sections"])))
    if analysis.get("bottom_sections"):
        frames.append(("أدنى_الشعب", pd.DataFrame(analysis["bottom_sections"])))
    if analysis.get("strongest_items") or analysis.get("weakest_items"):
        item_rows = []
        for q in analysis.get("strongest_items") or []:
            item_rows.append(
                {
                    "المحور": "قوة",
                    "البند": q.get("label_ar"),
                    "النسبة_%": q.get("score_percent"),
                    "التصنيف": q.get("classification_ar"),
                }
            )
        for q in analysis.get("weakest_items") or []:
            item_rows.append(
                {
                    "المحور": "تحسين",
                    "البند": q.get("label_ar"),
                    "النسبة_%": q.get("score_percent"),
                    "التصنيف": q.get("classification_ar"),
                }
            )
        frames.append(("بنود_القسم", pd.DataFrame(item_rows)))
    conclusion_rows = [{"#": i + 1, "الاستنتاج": c} for i, c in enumerate(analysis.get("conclusions") or [])]
    frames.append(("الاستنتاجات", pd.DataFrame(conclusion_rows or [{"#": 1, "الاستنتاج": "—"}])))
    rec_rows = analysis.get("recommendations") or []
    frames.append(
        (
            "التوصيات",
            pd.DataFrame(rec_rows)
            if rec_rows
            else pd.DataFrame(columns=["التوصية", "المسؤول", "الإطار"]),
        )
    )
    if analysis.get("follow_up_sections"):
        frames.append(("متابعة_الشعب", pd.DataFrame(analysis["follow_up_sections"])))
    return frames


def _department_approval_excel_rows() -> list[dict[str, str]]:
    """ورقة اعتماد القسم — للتوقيع اليدوي."""
    return [
        {"البند": "اعتماد القسم", "القيمة": ""},
        {"البند": "رأي رئيس القسم / ملاحظات إضافية", "القيمة": ""},
        {"البند": "", "القيمة": ""},
        {"البند": "", "القيمة": ""},
        {"البند": "التوقيع", "القيمة": ""},
        {"البند": "التاريخ", "القيمة": ""},
    ]


def _course_eval_sections_cover_rows(
    *,
    college_name_ar: str,
    department_label: str,
    semester: str,
    stats: dict[str, Any],
    course_eval_policy_ar: str,
    export_date: str,
    interpretation_ar: str | None = None,
) -> list[dict[str, str]]:
    avg = stats.get("avg_score_percent")
    avg_txt = f"{avg}%" if avg is not None else "—"
    rows = [
        {"البند": college_name_ar, "القيمة": ""},
        {"البند": "القسم", "القيمة": department_label},
        {"البند": "", "القيمة": ""},
        {"البند": "الوثيقة", "القيمة": "نتائج تقييم المقرر والأستاذ — حسب الشعبة"},
        {"البند": "الفصل الدراسي", "القيمة": semester},
        {"البند": "تاريخ التصدير", "القيمة": export_date},
        {"البند": "", "القيمة": ""},
        {"البند": "عدد الشعب", "القيمة": str(stats.get("section_count") or 0)},
        {"البند": "شعب مجمّعة", "القيمة": str(stats.get("aggregated_count") or 0)},
        {"البند": "شعب ناقصة", "القيمة": str(stats.get("pending_count") or 0)},
        {"البند": "متوسط النتائج المجمّعة", "القيمة": avg_txt},
        {"البند": "سياسة التجميع", "القيمة": course_eval_policy_ar},
    ]
    if interpretation_ar:
        rows.append({"البند": "التفسير العام", "القيمة": interpretation_ar})
    rows.append(
        {
            "البند": "الخصوصية",
            "القيمة": "لا تُعرض إجابات فردية — التجميع بعد بلوغ الحد الأدنى في الشعبة.",
        }
    )
    return rows


def _course_eval_section_evidence_rows(sections: list[dict]) -> list[dict]:
    """صفوف جدول الشعب للتصدير الرسمي — بلا ربط اعتماد."""
    rows: list[dict] = []
    for r in sections:
        rows.append(
            {
                "المقرر": r.get("course_name"),
                "الشعبة": r.get("section_id"),
                "الأستاذ": r.get("instructor_name"),
                "القسم": r.get("department_name"),
                "مسجّلون_تقدير": r.get("enrolled_count"),
                "إجمالي_تسجيل_المقرر": r.get("course_registration_count"),
                "عدد_التقييمات": r.get("response_count"),
                "نسبة_المشاركة_%": r.get("response_rate_percent"),
                "الحد_الأدنى": r.get("min_aggregate"),
                "حالة_التجميع": "مكتمل" if r.get("aggregated") else "ناقص",
                "النتيجة_%": r.get("overall_score_percent"),
                "أضعف_بند": r.get("weakest_item", "—"),
                "أقوى_بند": r.get("strongest_item", "—"),
            }
        )
    return rows


def _course_eval_by_course_evidence_rows(groups: list[dict]) -> list[dict]:
    rows: list[dict] = []
    for r in groups:
        rows.append(
            {
                "المقرر": r.get("course_name"),
                "الأستاذ": r.get("instructor_name"),
                "عدد_الشعب": r.get("section_count"),
                "معرّفات_الشعب": ", ".join(str(x) for x in (r.get("section_ids") or [])),
                "القسم": r.get("department_name"),
                "مسجّلون": r.get("enrolled_count"),
                "عدد_التقييمات": r.get("response_count"),
                "نسبة_المشاركة_%": r.get("response_rate_percent"),
                "الحد_الأدنى": r.get("min_aggregate"),
                "حالة_التجميع": "مكتمل" if r.get("aggregated") else "ناقص",
                "النتيجة_%": r.get("overall_score_percent"),
                "أضعف_بند": r.get("weakest_item", "—"),
                "أقوى_بند": r.get("strongest_item", "—"),
            }
        )
    return rows


def _course_eval_section_summary_rows(sections: list[dict]) -> list[dict]:
    rows: list[dict] = []
    for r in sections:
        rows.append(
            {
                "المقرر": r.get("course_name"),
                "الشعبة": r.get("section_id"),
                "الأستاذ": r.get("instructor_name"),
                "القسم": r.get("department_name"),
                "مسجّلون_تقدير": r.get("enrolled_count"),
                "إجمالي_تسجيل_المقرر": r.get("course_registration_count"),
                "عدد_التقييمات": r.get("response_count"),
                "نسبة_المشاركة_%": r.get("response_rate_percent"),
                "الحد_الأدنى": r.get("min_aggregate"),
                "حالة_التجميع": "مكتمل" if r.get("aggregated") else "ناقص",
                "النتيجة_%": r.get("overall_score_percent"),
                "أضعف_بند": r.get("weakest_item", "—"),
                "أقوى_بند": r.get("strongest_item", "—"),
                "معيار_الاعتماد": r.get("primary_accreditation"),
                "حالة_الامتثال": r.get("compliance_status_ar"),
            }
        )
    return rows


def _course_eval_section_detail_rows(sections: list[dict]) -> list[dict]:
    rows: list[dict] = []
    for r in sections:
        if not r.get("aggregated"):
            continue
        for q in r.get("questions") or []:
            rows.append(
                {
                    "المقرر": r.get("course_name"),
                    "الشعبة": r.get("section_id"),
                    "الأستاذ": r.get("instructor_name"),
                    "البند": q.get("label_ar"),
                    "متوسط_1_5": q.get("avg_rating"),
                    "النسبة_%": q.get("score_percent"),
                    "التصنيف": q.get("classification_ar"),
                }
            )
    return rows


def _course_eval_by_course_summary_rows(groups: list[dict]) -> list[dict]:
    rows: list[dict] = []
    for r in groups:
        rows.append(
            {
                "المقرر": r.get("course_name"),
                "الأستاذ": r.get("instructor_name"),
                "عدد_الشعب": r.get("section_count"),
                "معرّفات_الشعب": ", ".join(str(x) for x in (r.get("section_ids") or [])),
                "القسم": r.get("department_name"),
                "مسجّلون": r.get("enrolled_count"),
                "عدد_التقييمات": r.get("response_count"),
                "نسبة_المشاركة_%": r.get("response_rate_percent"),
                "الحد_الأدنى": r.get("min_aggregate"),
                "حالة_التجميع": "مكتمل" if r.get("aggregated") else "ناقص",
                "النتيجة_%": r.get("overall_score_percent"),
                "أضعف_بند": r.get("weakest_item", "—"),
                "أقوى_بند": r.get("strongest_item", "—"),
                "حالة_الامتثال": r.get("compliance_status_ar"),
            }
        )
    return rows


def package_excel_frames(
    combined: dict[str, Any],
) -> list[tuple[str, pd.DataFrame]]:
    """أوراق Excel للتقرير الموحّد."""
    reports = combined.get("reports") or []
    course_eval = combined.get("course_eval")
    course_sections = combined.get("course_eval_sections") or []
    course_by_course = combined.get("course_eval_by_course") or []
    frames: list[tuple[str, pd.DataFrame]] = [
        ("ملخص_تنفيذي", pd.DataFrame(_executive_summary_rows(reports, course_eval))),
        ("ربط_المعايير", pd.DataFrame(_accreditation_map_rows(reports, course_eval))),
        ("تحليل_مقارن", pd.DataFrame(_comparative_analysis_rows(reports, course_eval))),
    ]
    analysis = combined.get("analysis")
    if analysis:
        insert_at = 1
        for af_name, af_df in _analysis_excel_frames(analysis):
            frames.insert(insert_at, (af_name, af_df))
            insert_at += 1
    if course_sections:
        frames.append(
            ("ملخص_المقررات", pd.DataFrame(_course_eval_section_summary_rows(course_sections)))
        )
        detail_rows = _course_eval_section_detail_rows(course_sections)
        frames.append(
            (
                "بنود_المقررات",
                pd.DataFrame(detail_rows) if detail_rows else pd.DataFrame(columns=["المقرر"]),
            )
        )
    if course_by_course:
        frames.append(
            (
                "مقرر_وأستاذ",
                pd.DataFrame(_course_eval_by_course_summary_rows(course_by_course)),
            )
        )
    for r in reports:
        code = r.get("template_code") or ""
        frames.append((_sheet_name_for_code(code, r.get("title_ar") or ""), pd.DataFrame(_question_rows(r))))
    if course_eval:
        frames.append(
            (
                _sheet_name_for_code("student_course"),
                pd.DataFrame(_question_rows(course_eval)),
            )
        )
    frames.append(("بيانات_وصفية", pd.DataFrame(_metadata_rows(combined))))
    return frames


def course_eval_sections_excel_frames(
    sections: list[dict],
    *,
    by_course: list[dict] | None = None,
    semester: str = "",
    department_label: str = "",
    course_eval_policy_ar: str = "",
    missing_audit: dict[str, Any] | None = None,
    college_name_ar: str = COLLEGE_NAME_AR,
    export_date: str | None = None,
    analysis: dict[str, Any] | None = None,
) -> list[tuple[str, pd.DataFrame]]:
    """أوراق Excel لتصدير تقييم المقررات حسب الشعبة (شاهد رسمي)."""
    exp_date = export_date or datetime.datetime.now().strftime("%Y-%m-%d")
    stats = _course_eval_sections_export_stats(sections)
    frames: list[tuple[str, pd.DataFrame]] = [
        (
            "الغلاف",
            pd.DataFrame(
                _course_eval_sections_cover_rows(
                    college_name_ar=college_name_ar,
                    department_label=department_label,
                    semester=semester,
                    stats=stats,
                    course_eval_policy_ar=course_eval_policy_ar or format_course_eval_aggregation_policy(),
                    export_date=exp_date,
                    interpretation_ar=(analysis or {}).get("interpretation_ar"),
                )
            ),
        ),
    ]
    if analysis:
        frames.extend(_analysis_excel_frames(analysis))
    frames.append(("ملخص_الشعب", pd.DataFrame(_course_eval_section_evidence_rows(sections))))
    detail = _course_eval_section_detail_rows(sections)
    frames.append(
        (
            "بنود_الشعب",
            pd.DataFrame(detail) if detail else pd.DataFrame(columns=["المقرر", "البند"]),
        )
    )
    if by_course:
        frames.append(
            ("مقرر_وأستاذ", pd.DataFrame(_course_eval_by_course_evidence_rows(by_course)))
        )
    if missing_audit and (missing_audit.get("rows") or []):
        miss_frames = course_eval_missing_audit_excel_frames(missing_audit)
        for sheet_name, df in miss_frames:
            if sheet_name == "شعب_بلا_تقييم":
                frames.append((sheet_name, df))
                break
    meta = [
        {"البند": "الفصل", "القيمة": semester},
        {"البند": "القسم", "القيمة": department_label},
        {"البند": "عدد الشعب", "القيمة": len(sections)},
        {"البند": "شعب مجمّعة", "القيمة": stats.get("aggregated_count")},
        {"البند": "شعب ناقصة", "القيمة": stats.get("pending_count")},
        {
            "البند": "سياسة التجميع",
            "القيمة": course_eval_policy_ar or format_course_eval_aggregation_policy(),
        },
        {
            "البند": "منهجية القياس",
            "القيمة": "مقياس Likert 1–5؛ النسبة = (متوسط البند / 5) × 100",
        },
        {"البند": "تاريخ التصدير", "القيمة": exp_date},
    ]
    frames.append(("بيانات_وصفية", pd.DataFrame(meta)))
    frames.append(("اعتماد_القسم", pd.DataFrame(_department_approval_excel_rows())))
    return frames


def course_eval_sections_excel_bytes(
    sections: list[dict],
    *,
    by_course: list[dict] | None = None,
    semester: str = "",
    department_label: str = "",
    course_eval_policy_ar: str = "",
    missing_audit: dict[str, Any] | None = None,
    college_name_ar: str = COLLEGE_NAME_AR,
    export_date: str | None = None,
    analysis: dict[str, Any] | None = None,
) -> bytes:
    """بايتات Excel مع تنسيق عربي RTL."""
    frames = course_eval_sections_excel_frames(
        sections,
        by_course=by_course,
        semester=semester,
        department_label=department_label,
        course_eval_policy_ar=course_eval_policy_ar,
        missing_audit=missing_audit,
        college_name_ar=college_name_ar,
        export_date=export_date,
        analysis=analysis,
    )
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        workbook = writer.book
        formats = excel_arabic_workbook_formats(workbook)
        used_names: set[str] = set()
        for sheet_name, df in frames:
            from backend.services.utilities import _sanitize_excel_sheet_name

            safe_name = _sanitize_excel_sheet_name(sheet_name, used_names)
            data = df if isinstance(df, pd.DataFrame) else pd.DataFrame(df)
            data.to_excel(writer, index=False, sheet_name=safe_name)
            ws = writer.sheets[safe_name]
            write_excel_sheet_rtl(
                ws,
                data,
                formats=formats,
                cover_sheet=safe_name in ("الغلاف", "اعتماد_القسم"),
            )
    buf.seek(0)
    return buf.getvalue()


def build_course_eval_sections_export_context(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
) -> dict[str, Any]:
    """سياق مشترك لتصدير تقييم المقررات حسب الشعبة (Excel / Word / PDF)."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    sections, by_course = build_course_eval_results_bundle(
        conn, semester=sem, department_id=department_id
    )
    missing_audit = build_course_eval_missing_sections_audit(
        conn, semester=sem, department_id=department_id
    )
    policy = format_course_eval_aggregation_policy(conn)
    dept_label = _department_label(conn, department_id)
    export_date = datetime.datetime.now().strftime("%Y-%m-%d")
    stats = _course_eval_sections_export_stats(sections)
    sem_slug = sem.replace(" ", "_")[:40]
    analysis = build_course_eval_sections_analysis(
        sections=sections,
        stats=stats,
        missing_audit=missing_audit,
        department_label=dept_label,
        semester=sem,
        course_eval_policy_ar=policy,
    )
    return {
        "college_name_ar": COLLEGE_NAME_AR,
        "department_label": dept_label,
        "semester": sem,
        "export_date": export_date,
        "sections": sections,
        "by_course": by_course,
        "missing_audit": missing_audit,
        "course_eval_policy_ar": policy,
        "stats": stats,
        "analysis": analysis,
        "title": "نتائج تقييم المقرر والأستاذ — حسب الشعبة",
        "filename_prefix": f"course_eval_sections_{sem_slug}",
        "pdf_arabic_css": pdf_arabic_extra_css(for_pdf=False),
        "pdf_arabic_css_print": pdf_arabic_extra_css(for_pdf=True),
        "methodology_rows": [
            {"البند": "مقياس التقييم", "القيمة": "Likert 1–5"},
            {"البند": "طريقة الحساب", "القيمة": "النسبة = (متوسط البند / 5) × 100"},
            {"البند": "قاعدة التجميع", "القيمة": policy},
            {"البند": "الخصوصية", "القيمة": "لا تُعرض إجابات فردية في التصدير"},
        ],
    }


def course_eval_sections_docx_bytes(ctx: dict[str, Any]) -> bytes:
    """مستند Word لتقرير الشعب مع تحليل واستنتاجات."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as exc:
        raise RuntimeError(
            "مكتبة python-docx غير متوفرة. ثبّت docxtpl أو python-docx."
        ) from exc

    doc = Document()
    college = ctx.get("college_name_ar") or COLLEGE_NAME_AR
    dept = ctx.get("department_label") or "—"
    sem = ctx.get("semester") or "—"
    stats = ctx.get("stats") or {}
    policy = ctx.get("course_eval_policy_ar") or ""
    analysis = ctx.get("analysis") or {}

    docx_add_rtl_paragraph(doc, college, bold=True, center=True, font_size=16)
    docx_add_rtl_paragraph(doc, f"القسم: {dept}", bold=True, center=True, font_size=14)
    docx_add_rtl_paragraph(
        doc,
        ctx.get("title") or "نتائج تقييم المقرر والأستاذ — حسب الشعبة",
        bold=True,
        center=True,
        font_size=13,
    )
    docx_add_rtl_paragraph(doc, f"الفصل الدراسي: {sem}")
    docx_add_rtl_paragraph(doc, f"تاريخ التصدير: {ctx.get('export_date') or '—'}")

    docx_add_rtl_heading(doc, "ملخص تنفيذي", level=2)
    avg = stats.get("avg_score_percent")
    avg_txt = f"{avg}%" if avg is not None else "—"
    docx_add_rtl_paragraph(
        doc,
        f"عدد الشعب: {stats.get('section_count') or 0} — "
        f"مجمّعة: {stats.get('aggregated_count') or 0} — "
        f"ناقصة: {stats.get('pending_count') or 0} — "
        f"متوسط النتائج المجمّعة: {avg_txt}",
    )
    docx_add_rtl_paragraph(doc, f"سياسة التجميع: {policy}")
    if analysis.get("interpretation_ar"):
        docx_add_rtl_paragraph(doc, analysis["interpretation_ar"], italic=True)

    docx_add_rtl_heading(doc, "التحليل", level=2)
    for para in analysis.get("narrative_paragraphs") or []:
        docx_add_rtl_paragraph(doc, para)

    if analysis.get("distribution_rows"):
        docx_add_rtl_heading(doc, "توزيع النتائج", level=3)
        dist_table = doc.add_table(rows=1, cols=3)
        dist_table.style = "Table Grid"
        docx_fill_rtl_table(
            dist_table,
            ["التصنيف", "عدد الشعب", "أمثلة"],
            [
                [r.get("التصنيف"), r.get("عدد_الشعب"), r.get("أمثلة")]
                for r in analysis["distribution_rows"]
            ],
        )

    top_secs = analysis.get("top_sections") or []
    bottom_secs = analysis.get("bottom_sections") or []
    if top_secs or bottom_secs:
        docx_add_rtl_heading(doc, "أعلى وأدنى الشعب", level=3)
        rank_table = doc.add_table(rows=1, cols=5)
        rank_table.style = "Table Grid"
        rank_rows = []
        for r in top_secs:
            rank_rows.append(
                [
                    "أعلى",
                    r.get("المقرر"),
                    r.get("الشعبة"),
                    r.get("الأستاذ"),
                    f"{r.get('النتيجة_%')}%",
                ]
            )
        for r in bottom_secs:
            rank_rows.append(
                [
                    "أدنى",
                    r.get("المقرر"),
                    r.get("الشعبة"),
                    r.get("الأستاذ"),
                    f"{r.get('النتيجة_%')}%",
                ]
            )
        docx_fill_rtl_table(
            rank_table,
            ["التصنيف", "المقرر", "الشعبة", "الأستاذ", "النتيجة %"],
            rank_rows,
        )

    strongest = analysis.get("strongest_items") or []
    weakest = analysis.get("weakest_items") or []
    if strongest or weakest:
        docx_add_rtl_heading(doc, "ملحق ج — أبرز البنود على مستوى القسم", level=3)
        item_table = doc.add_table(rows=1, cols=4)
        item_table.style = "Table Grid"
        item_rows = []
        for q in strongest:
            item_rows.append(["قوة", q.get("label_ar"), f"{q.get('score_percent')}%", q.get("classification_ar")])
        for q in weakest:
            item_rows.append(["تحسين", q.get("label_ar"), f"{q.get('score_percent')}%", q.get("classification_ar")])
        docx_fill_rtl_table(item_table, ["المحور", "البند", "النسبة", "التصنيف"], item_rows)

    docx_add_rtl_heading(doc, "الاستنتاجات", level=2)
    for i, c in enumerate(analysis.get("conclusions") or [], start=1):
        docx_add_rtl_paragraph(doc, f"{i}. {c}")

    recs = analysis.get("recommendations") or []
    if recs:
        docx_add_rtl_heading(doc, "التوصيات", level=2)
        rec_table = doc.add_table(rows=1, cols=3)
        rec_table.style = "Table Grid"
        docx_fill_rtl_table(
            rec_table,
            ["التوصية", "المسؤول", "الإطار"],
            [[r.get("التوصية"), r.get("المسؤول"), r.get("الإطار")] for r in recs],
        )

    docx_add_rtl_heading(doc, "ملحق أ — جدول الشعب", level=2)
    headers = [
        "المقرر",
        "الشعبة",
        "الأستاذ",
        "القسم",
        "مسجّلون",
        "تقييمات",
        "الحد",
        "النتيجة %",
        "التجميع",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    section_rows = []
    for s in ctx.get("sections") or []:
        pct = s.get("overall_score_percent")
        section_rows.append(
            [
                s.get("course_name"),
                s.get("section_id"),
                s.get("instructor_name"),
                s.get("department_name"),
                s.get("enrolled_count") if s.get("enrolled_count") is not None else "—",
                s.get("response_count") if s.get("response_count") is not None else "—",
                s.get("min_aggregate") if s.get("min_aggregate") is not None else "—",
                f"{pct}%" if pct is not None and s.get("aggregated") else "—",
                "مكتمل" if s.get("aggregated") else "ناقص",
            ]
        )
    docx_fill_rtl_table(table, headers, section_rows)

    missing = (ctx.get("missing_audit") or {}).get("rows") or []
    if missing:
        docx_add_rtl_heading(doc, "ملحق ب — شعب بلا تقييم", level=2)
        mtable = doc.add_table(rows=1, cols=6)
        mtable.style = "Table Grid"
        docx_fill_rtl_table(
            mtable,
            ["المقرر", "الشعبة", "الأستاذ", "القسم", "مسجّلون", "أسباب الفجوة"],
            [
                [
                    r.get("course_name"),
                    r.get("section_id"),
                    r.get("instructor_name"),
                    r.get("department_name"),
                    r.get("enrolled_count") if r.get("enrolled_count") is not None else "—",
                    r.get("gap_reasons_ar") or "—",
                ]
                for r in missing
            ],
        )

    follow = analysis.get("follow_up_sections") or []
    if follow:
        docx_add_rtl_heading(doc, "ملحق د — شعب تحتاج متابعة", level=2)
        ftable = doc.add_table(rows=1, cols=4)
        ftable.style = "Table Grid"
        docx_fill_rtl_table(
            ftable,
            ["المقرر", "الشعبة", "الأستاذ", "السبب"],
            [
                [r.get("المقرر"), r.get("الشعبة"), r.get("الأستاذ"), r.get("السبب")]
                for r in follow
            ],
        )

    docx_add_rtl_heading(doc, "المنهجية والخصوصية", level=2)
    for row in ctx.get("methodology_rows") or []:
        docx_add_rtl_paragraph(doc, f"{row.get('البند')}: {row.get('القيمة')}")

    docx_add_rtl_heading(doc, "اعتماد القسم", level=2)
    docx_add_rtl_paragraph(doc, "رأي رئيس القسم / ملاحظات إضافية:")
    docx_add_rtl_paragraph(doc, "_" * 60)
    docx_add_rtl_paragraph(doc, "التوقيع: ___________________     التاريخ: ___________________")

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_docx_paragraph_rtl(foot, align_right=False)
    fr = foot.add_run(
        "مُنشأ آلياً من منصة ضمان الجودة — لا يُعتمد دون توقيع الجهة المختصة عند الحاجة."
    )
    fr.italic = True
    set_docx_run_arabic_font(fr)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = tmp.name
    tmp.close()
    doc.save(tmp_path)
    with open(tmp_path, "rb") as fh:
        raw = fh.read()
    try:
        import os

        os.unlink(tmp_path)
    except OSError:
        pass
    return raw


def single_survey_excel_frames(report: dict[str, Any]) -> list[tuple[str, pd.DataFrame]]:
    """أوراق Excel لتقرير استبيان واحد."""
    summary = {
        "الاستبيان": report.get("title_ar"),
        "الرمز": report.get("template_code"),
        "الفصل": report.get("semester"),
        "النطاق": report.get("department_label") or report.get("department_name"),
        "الفئة": report.get("respondent_label"),
        "عدد_الإجابات": report.get("response_count"),
        "الحد_الأدنى": report.get("min_aggregate"),
        "حالة_التجميع": "مكتمل" if report.get("aggregated") else "ناقص",
        "النتيجة_%": report.get("overall_score_percent"),
        "حالة_الامتثال": report.get("compliance_status_ar"),
    }
    if report.get("course_name"):
        summary["المقرر"] = report.get("course_name")
    if report.get("instructor_name"):
        summary["الأستاذ"] = report.get("instructor_name")
    if report.get("section_id"):
        summary["الشعبة"] = report.get("section_id")
    if report.get("section_count") and int(report.get("section_count") or 0) > 1:
        summary["عدد_الشعب"] = report.get("section_count")
        summary["معرّفات_الشعب"] = ", ".join(
            str(x) for x in (report.get("section_ids") or [])
        )
    acc_rows = [
        {
            "المعيار": l.get("standard_code"),
            "المؤشر": l.get("indicator_code"),
            "عنوان_المؤشر": l.get("indicator_title_ar"),
            "نوع_الربط": LINK_TYPE_LABELS_AR.get(l.get("link_type", ""), l.get("link_type")),
            "كيفية_الاستفادة": l.get("usage_ar"),
        }
        for l in report.get("accreditation_links") or []
    ]
    rec_rows = [{"التوصية": r} for r in report.get("recommendations") or []]
    method_rows = [
        {"البند": "مقياس التقييم", "القيمة": "Likert 1–5"},
        {"البند": "طريقة الحساب", "القيمة": "النسبة = (متوسط البند / 5) × 100"},
        {"البند": "مسجّلون (تقدير)", "القيمة": report.get("enrolled_count")},
        {"البند": "الحد الأدنى للتجميع", "القيمة": report.get("min_aggregate")},
        {
            "البند": "قاعدة التجميع",
            "القيمة": report.get("course_eval_policy_ar") or format_course_eval_aggregation_policy(),
        },
        {"البند": "الخصوصية", "القيمة": "لا تُعرض إجابات فردية في التصدير"},
    ]
    analysis = report.get("analysis") or build_survey_report_analysis(report)
    frames: list[tuple[str, pd.DataFrame]] = [
        ("ملخص", pd.DataFrame([summary])),
    ]
    frames.extend(_analysis_excel_frames(analysis))
    frames.extend(
        [
        ("البنود", pd.DataFrame(_question_rows(report))),
        ("المعايير", pd.DataFrame(acc_rows) if acc_rows else pd.DataFrame(columns=["المؤشر"])),
        ("توصيات", pd.DataFrame(rec_rows) if rec_rows else pd.DataFrame(columns=["التوصية"])),
        ("منهجية", pd.DataFrame(method_rows)),
        ]
    )
    return frames


def export_package_xlsx(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
    include_course_eval: bool = True,
):
    combined = build_combined_survey_report(
        conn,
        semester=semester,
        department_id=department_id,
        include_course_eval=include_course_eval,
    )
    combined["analysis"] = build_combined_survey_analysis(combined)
    from backend.services.survey_report_charts import build_chart_data_for_combined

    chart_data = build_chart_data_for_combined(combined, combined["analysis"])
    sem_slug = (combined.get("semester") or "report").replace(" ", "_")[:40]
    now = datetime.datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    fname = f"survey_package_{sem_slug}_{now}.xlsx"
    raw = survey_excel_bytes_from_frames(package_excel_frames(combined), chart_data=chart_data)
    return send_file(
        io.BytesIO(raw),
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=fname,
    )


def export_single_survey_xlsx(
    conn,
    template_code: str,
    *,
    semester: str | None = None,
    department_id: int | None = None,
):
    code = (template_code or "").strip()
    if code == "student_course":
        report = build_course_eval_report(conn, semester=semester, department_id=department_id)
        report["department_label"] = _department_label(conn, department_id)
        report = _enrich_course_eval_report_for_display(report)
    else:
        report = build_survey_report(conn, code, semester=semester, department_id=department_id)
    report["analysis"] = build_survey_report_analysis(report)
    from backend.services.survey_report_charts import build_chart_data_for_survey

    chart_data = build_chart_data_for_survey(report, report["analysis"])
    sem_slug = (report.get("semester") or "report").replace(" ", "_")[:40]
    now = datetime.datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    fname = f"survey_{code}_{sem_slug}_{now}.xlsx"
    raw = survey_excel_bytes_from_frames(
        single_survey_excel_frames(report), chart_data=chart_data
    )
    return send_file(
        io.BytesIO(raw),
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=fname,
    )


def is_exportable_template_code(conn, template_code: str) -> bool:
    code = (template_code or "").strip()
    if code in ("student_course", "course_eval_sections"):
        return True
    return get_template_by_code(conn, code) is not None


def export_course_eval_sections_xlsx(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
):
    ctx = build_course_eval_sections_export_context(
        conn, semester=semester, department_id=department_id
    )
    now = datetime.datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    fname = f"{ctx['filename_prefix']}_{now}.xlsx"
    raw = course_eval_sections_excel_bytes(
        ctx["sections"],
        by_course=ctx.get("by_course"),
        semester=ctx["semester"],
        department_label=ctx["department_label"],
        course_eval_policy_ar=ctx.get("course_eval_policy_ar"),
        missing_audit=ctx.get("missing_audit"),
        college_name_ar=ctx.get("college_name_ar"),
        export_date=ctx.get("export_date"),
        analysis=ctx.get("analysis"),
    )
    return send_file(
        io.BytesIO(raw),
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=True,
        download_name=fname,
    )


def export_course_eval_sections_docx(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
):
    ctx = build_course_eval_sections_export_context(
        conn, semester=semester, department_id=department_id
    )
    raw = course_eval_sections_docx_bytes(ctx)
    now = datetime.datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    fname = f"{ctx['filename_prefix']}_{now}.docx"
    return send_file(
        io.BytesIO(raw),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=fname,
    )


def course_eval_sections_export_bytes(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
    fmt: str = "xlsx",
) -> tuple[bytes, str, dict[str, Any]]:
    """بايتات التصدير لرفع الشاهد — fmt: xlsx | docx."""
    ctx = build_course_eval_sections_export_context(
        conn, semester=semester, department_id=department_id
    )
    sem_slug = (ctx.get("semester") or "report").replace(" ", "_")[:40]
    export_fmt = (fmt or "xlsx").strip().lower()
    if export_fmt in ("docx", "word"):
        raw = course_eval_sections_docx_bytes(ctx)
        filename = f"course_eval_sections_{sem_slug}.docx"
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        raw = course_eval_sections_excel_bytes(
            ctx["sections"],
            by_course=ctx.get("by_course"),
            semester=ctx["semester"],
            department_label=ctx["department_label"],
            course_eval_policy_ar=ctx.get("course_eval_policy_ar"),
            missing_audit=ctx.get("missing_audit"),
            college_name_ar=ctx.get("college_name_ar"),
            export_date=ctx.get("export_date"),
            analysis=ctx.get("analysis"),
        )
        filename = f"course_eval_sections_{sem_slug}.xlsx"
        mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    report = {
        "title_ar": ctx.get("title"),
        "semester": ctx.get("semester"),
        "department_label": ctx.get("department_label"),
        "response_count": ctx.get("stats", {}).get("aggregated_count"),
        "overall_score_percent": ctx.get("stats", {}).get("avg_score_percent"),
        "_mime": mime,
    }
    return raw, filename, report


def export_course_eval_section_xlsx(
    conn,
    section_id: int,
    *,
    semester: str | None = None,
):
    sem = (semester or "").strip() or term_label_from_conn(conn)
    report = build_course_eval_section_report(conn, int(section_id), semester=sem)
    if not report:
        return None
    sem_slug = sem.replace(" ", "_")[:40]
    return excel_response_from_frames(
        single_survey_excel_frames(report),
        filename_prefix=f"course_eval_section_{int(section_id)}_{sem_slug}",
    )


def export_course_eval_by_course_xlsx(
    conn,
    course_name: str,
    instructor_id: int,
    *,
    semester: str | None = None,
    department_id: int | None = None,
):
    sem = (semester or "").strip() or term_label_from_conn(conn)
    report = build_course_eval_by_course_report(
        conn,
        course_name,
        int(instructor_id),
        semester=sem,
        department_id=department_id,
    )
    if not report:
        return None
    sem_slug = sem.replace(" ", "_")[:40]
    safe_course = (course_name or "course").replace(" ", "_")[:30]
    return excel_response_from_frames(
        single_survey_excel_frames(report),
        filename_prefix=f"course_eval_{safe_course}_inst{int(instructor_id)}_{sem_slug}",
    )


OVERALL_INTERPRETATION_AR: dict[str, str] = {
    "excellent": "أداء ممتاز يدعم متطلبات الاعتماد — يُوصى بتوثيق الممارسات كشواهد.",
    "good": "أداء جيد بشكل عام — مراقبة دورية في اجتماعات ضمان الجودة.",
    "needs_improvement": "نتائج دون المستوى المطلوب — خطة تحسين فصلية مطلوبة.",
    "critical": "نتائج حرجة — إجراء تصحيحي عاجل مع المعنيين.",
    "pending": "لا تتوفر نتيجة مجمّعة — عدد الإجابات دون الحد الأدنى المطلوب للخصوصية.",
}


def interpret_overall_score_ar(percent: float | None, aggregated: bool) -> str:
    if not aggregated or percent is None:
        return OVERALL_INTERPRETATION_AR["pending"]
    cls = classify_item_score(percent)
    return OVERALL_INTERPRETATION_AR.get(cls, "—")


def _enrich_course_eval_report_for_display(report: dict[str, Any]) -> dict[str, Any]:
    questions = _enrich_questions(report.get("questions") or [])
    weakest, strongest = _weakest_strongest(questions)
    title = report.get("title_ar") or "تقييم المقرر"
    return {
        **report,
        "questions": questions,
        "weakest_item": weakest,
        "strongest_item": strongest,
        "recommendations": generate_recommendations(questions, title),
        "interpretation_ar": interpret_overall_score_ar(
            report.get("overall_score_percent"), report.get("aggregated")
        ),
    }


def build_survey_report_analysis(report: dict[str, Any]) -> dict[str, Any]:
    """تحليل واستنتاجات آلية لاستبيان واحد (شاهد / معاينة)."""
    from backend.services.survey_report_charts import (
        ITEM_CLASS_DISTRIBUTION_LABELS,
        ITEM_CLASS_BUCKET_ORDER,
        item_distribution_buckets,
    )

    title = (report.get("title_ar") or report.get("template_code") or "الاستبيان").strip()
    questions = report.get("questions") or []
    aggregated = bool(report.get("aggregated"))
    score = report.get("overall_score_percent")
    interpretation = report.get("interpretation_ar") or interpret_overall_score_ar(score, aggregated)
    dept = report.get("department_label") or report.get("department_name") or "—"
    sem = report.get("semester") or report.get("cycle_label") or "—"
    resp_n = int(report.get("response_count") or 0)
    min_agg = report.get("min_aggregate")

    buckets = item_distribution_buckets(questions)
    distribution_rows = [
        {
            "التصنيف": ITEM_CLASS_DISTRIBUTION_LABELS[key],
            "العدد": len(bucket),
            "أمثلة": "؛ ".join(
                _truncate_item_label(q.get("label_ar") or "") for q in bucket[:3]
            )
            or "—",
        }
        for key in ITEM_CLASS_BUCKET_ORDER
        for bucket in [buckets[key]]
        if bucket
    ]

    scored = [q for q in questions if q.get("score_percent") is not None]
    sorted_q = sorted(scored, key=lambda x: float(x["score_percent"]), reverse=True)
    strongest_items = [
        {
            "label_ar": q.get("label_ar"),
            "score_percent": q.get("score_percent"),
            "classification_ar": q.get("classification_ar"),
        }
        for q in sorted_q[:3]
    ]
    weakest_items = [
        {
            "label_ar": q.get("label_ar"),
            "score_percent": q.get("score_percent"),
            "classification_ar": q.get("classification_ar"),
        }
        for q in (list(reversed(sorted_q[-3:])) if sorted_q else [])
    ]

    narrative: list[str] = [
        (
            f"يغطي هذا التقرير استبيان «{title}» للفصل/الدورة «{sem}» ضمن نطاق «{dept}». "
            f"عدد الإجابات: {resp_n} (الحد الأدنى للتجميع: {min_agg or '—'})."
        ),
        (
            f"النتيجة الإجمالية: {score}% — {interpretation}"
            if aggregated and score is not None
            else "لم يكتمل التجميع — لا تُعرض بنود تفصيلية حتى بلوغ الحد الأدنى."
        ),
    ]
    if strongest_items:
        labels = "، ".join(f"«{q['label_ar']}» ({q['score_percent']}%)" for q in strongest_items)
        narrative.append(f"أبرز نقاط القوة: {labels}.")
    if weakest_items:
        labels = "، ".join(f"«{q['label_ar']}» ({q['score_percent']}%)" for q in weakest_items)
        narrative.append(f"أبرز محاور التحسين: {labels}.")

    conclusions: list[str] = []
    if aggregated:
        conclusions.append("اكتمل التجميع الإحصائي وفق سياسة الخصوصية.")
    else:
        conclusions.append("التجميع ناقص — يُعاد التقرير بعد بلوغ الحد الأدنى للإجابات.")
    if score is not None and score >= 80:
        conclusions.append("النتيجة الإجمالية إيجابية وتدعم متطلبات ضمان الجودة.")
    elif score is not None and score >= 70:
        conclusions.append("النتيجة جيدة — يُنصح بمتابعة البنود الأضعف.")
    elif score is not None:
        conclusions.append("النتيجة دون المستوى المطلوب — خطة تحسين فصلية مطلوبة.")
    if buckets.get("critical"):
        conclusions.append(f"وُجد {len(buckets['critical'])} بند(اً) بحالة «حرج» — أولوية معالجة.")
    if not conclusions:
        conclusions.append("لا تتوفر بيانات كافية لاستنتاجات نهائية.")

    recommendations: list[dict[str, str]] = []
    for q in weakest_items[:2]:
        recommendations.append(
            {
                "التوصية": f"معالجة بند «{q['label_ar']}» ({q['score_percent']}%) في اجتماع لجنة الجودة.",
                "المسؤول": "رئيس القسم / لجنة الجودة",
                "الإطار": "خلال الفصل الحالي",
            }
        )
    for q in buckets.get("critical", [])[:2]:
        recommendations.append(
            {
                "التوصية": f"إجراء تصحيحي عاجل لبند «{q.get('label_ar')}» ({q.get('score_percent')}%).",
                "المسؤول": "الإدارة المعنية",
                "الإطار": "خلال 30 يوماً",
            }
        )
    if not recommendations and aggregated:
        recommendations.append(
            {
                "التوصية": f"توثيق ممارسات «{title}» كشاهد اعتماد عند استمرار النتائج الإيجابية.",
                "المسؤول": "لجنة الجودة",
                "الإطار": "نهاية الفصل",
            }
        )
    if not recommendations:
        recommendations.append(
            {
                "التوصية": "متابعة التعبئة حتى اكتمال الحد الأدنى ثم إعادة التقييم.",
                "المسؤول": "رئيس القسم",
                "الإطار": "أسبوعان",
            }
        )

    return {
        "interpretation_ar": interpretation,
        "narrative_paragraphs": narrative,
        "distribution_rows": distribution_rows,
        "strongest_items": strongest_items,
        "weakest_items": weakest_items,
        "conclusions": conclusions,
        "recommendations": recommendations,
    }


def _truncate_item_label(text: str, max_len: int = 48) -> str:
    s = str(text or "").strip()
    if len(s) <= max_len:
        return s
    return s[: max_len - 1] + "…"


def build_combined_survey_analysis(combined: dict[str, Any]) -> dict[str, Any]:
    """تحليل مقارن للتقرير الموحّد."""
    from backend.services.survey_report_charts import (
        ITEM_CLASS_DISTRIBUTION_LABELS,
        ITEM_CLASS_BUCKET_ORDER,
    )

    reports = combined.get("reports") or []
    aggregated = [r for r in reports if r.get("aggregated") and r.get("overall_score_percent") is not None]
    sem = combined.get("semester") or combined.get("cycle_label") or "—"
    dept = combined.get("department_label") or "—"
    total = int(combined.get("total_survey_count") or len(reports))
    agg_n = int(combined.get("aggregated_survey_count") or len(aggregated))

    buckets: dict[str, list[dict]] = {k: [] for k in ITEM_CLASS_BUCKET_ORDER}
    for r in aggregated:
        cls = classify_item_score(float(r["overall_score_percent"]))
        buckets[cls].append(r)

    distribution_rows = [
        {
            "التصنيف": ITEM_CLASS_DISTRIBUTION_LABELS[key],
            "العدد": len(bucket),
            "أمثلة": "؛ ".join(
                _truncate_item_label(r.get("title_ar") or r.get("template_code") or "") for r in bucket[:3]
            )
            or "—",
        }
        for key in ITEM_CLASS_BUCKET_ORDER
        for bucket in [buckets[key]]
        if bucket
    ]

    sorted_agg = sorted(aggregated, key=lambda x: float(x["overall_score_percent"]), reverse=True)
    top_surveys = [
        {
            "الاستبيان": r.get("title_ar"),
            "الفئة": r.get("respondent_label"),
            "النتيجة_%": r.get("overall_score_percent"),
        }
        for r in sorted_agg[:3]
    ]
    bottom_surveys = [
        {
            "الاستبيان": r.get("title_ar"),
            "الفئة": r.get("respondent_label"),
            "النتيجة_%": r.get("overall_score_percent"),
        }
        for r in sorted(sorted_agg, key=lambda x: float(x["overall_score_percent"]))[:3]
    ]

    narrative = list(generate_executive_narrative_ar(combined))
    role_avgs: dict[str, list[float]] = {}
    for r in aggregated:
        role = (r.get("respondent_label") or "—").strip()
        role_avgs.setdefault(role, []).append(float(r["overall_score_percent"]))
    if role_avgs:
        parts = [
            f"«{role}»: {round(sum(v)/len(v), 1)}%"
            for role, v in sorted(role_avgs.items())
        ]
        narrative.append("متوسط النتائج حسب فئة المستجيب: " + "؛ ".join(parts) + ".")

    conclusions: list[str] = []
    if agg_n >= max(1, total // 2):
        conclusions.append("تغطية التجميع على مستوى الاستبيانات مقبولة بشكل عام.")
    elif total:
        conclusions.append("عدد الاستبيانات المجمّعة دون المطلوب — يُعزّز حث المستجيبين.")
    if sorted_agg and float(sorted_agg[0]["overall_score_percent"]) >= 80:
        conclusions.append("هناك استبيانات بأداء ممتاز يمكن توثيقها كشواهد.")
    if buckets.get("critical") or buckets.get("needs_improvement"):
        conclusions.append("بعض الاستبيانات تستدعي خطط تحسين — راجع الجدول التفصيلي.")
    if not conclusions:
        conclusions.append("لا تتوفر بيانات كافية لاستنتاجات شاملة.")

    recommendations: list[dict[str, str]] = []
    for r in bottom_surveys[:2]:
        recommendations.append(
            {
                "التوصية": (
                    f"مراجعة استبيان «{r.get('الاستبيان')}» "
                    f"({r.get('النتيجة_%')}%) ووضع خطة تحسين."
                ),
                "المسؤول": "لجنة الجودة",
                "الإطار": "خلال الفصل الحالي",
            }
        )
    if not recommendations:
        recommendations.append(
            {
                "التوصية": "مناقشة نتائج الاستبيانات في اجتماع ضمان الجودة الدوري.",
                "المسؤول": "رئيس القسم",
                "الإطار": "نهاية الفصل",
            }
        )

    return {
        "narrative_paragraphs": narrative,
        "distribution_rows": distribution_rows,
        "top_sections": top_surveys,
        "bottom_sections": bottom_surveys,
        "conclusions": conclusions,
        "recommendations": recommendations,
    }


def enrich_survey_export_context(ctx: dict[str, Any], *, for_pdf: bool = False) -> dict[str, Any]:
    """إثراء سياق التصدير بالتحليل والرسوم."""
    import logging

    from backend.services.survey_report_charts import (
        build_chart_data_for_combined,
        build_chart_data_for_survey,
        build_chart_images_for_combined,
        build_chart_images_for_survey,
    )

    log = logging.getLogger(__name__)

    ctx["pdf_arabic_css"] = pdf_arabic_extra_css(for_pdf=False)
    ctx["pdf_arabic_css_print"] = pdf_arabic_extra_css(for_pdf=True)
    ctx.setdefault("college_name_ar", COLLEGE_NAME_AR)
    ctx.setdefault("export_date", datetime.datetime.now().strftime("%Y-%m-%d"))

    if ctx.get("report"):
        report_obj = ctx["report"]
        if not ctx.get("analysis") and not report_obj.get("has_segment_detail"):
            ctx["analysis"] = build_survey_report_analysis(report_obj)
        if ctx.get("analysis"):
            ctx["chart_data"] = build_chart_data_for_survey(report_obj, ctx.get("analysis"))
            if for_pdf:
                try:
                    ctx["chart_images"] = build_chart_images_for_survey(report_obj, ctx.get("analysis"))
                except Exception:
                    log.exception("survey chart_images (single) skipped")
                    ctx["chart_images"] = {}
    elif ctx.get("reports") is not None:
        if not ctx.get("analysis"):
            ctx["analysis"] = build_combined_survey_analysis(ctx)
        ctx["chart_data"] = build_chart_data_for_combined(ctx, ctx.get("analysis"))
        if for_pdf:
            try:
                ctx["chart_images"] = build_chart_images_for_combined(ctx, ctx.get("analysis"))
            except Exception:
                log.exception("survey chart_images (combined) skipped")
                ctx["chart_images"] = {}
    return ctx


def survey_excel_bytes_from_frames(
    frames: list[tuple[str, pd.DataFrame]],
    *,
    chart_data: dict[str, Any] | None = None,
    rtl: bool = True,
) -> bytes:
    """بايتات Excel مع تنسيق عربي وورقة رسوم اختيارية."""
    from backend.services.utilities import _sanitize_excel_sheet_name
    from backend.services.survey_report_charts import add_chart_sheet_to_workbook

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="xlsxwriter") as writer:
        workbook = writer.book
        formats = excel_arabic_workbook_formats(workbook) if rtl else {}
        used_names: set[str] = set()
        for sheet_name, df in frames:
            safe_name = _sanitize_excel_sheet_name(sheet_name, used_names)
            data = df if isinstance(df, pd.DataFrame) else pd.DataFrame(df)
            data.to_excel(writer, index=False, sheet_name=safe_name)
            if rtl and formats:
                ws = writer.sheets[safe_name]
                write_excel_sheet_rtl(
                    ws,
                    data,
                    formats=formats,
                    cover_sheet=safe_name in ("الغلاف", "اعتماد_القسم", "ملخص"),
                )
        if chart_data and chart_data.get("has_data"):
            add_chart_sheet_to_workbook(workbook, chart_data)
    buf.seek(0)
    return buf.getvalue()


def generate_executive_narrative_ar(combined: dict[str, Any]) -> list[str]:
    """فقرات تفسيرية آلية للملخص التنفيذي في PDF."""
    paragraphs: list[str] = []
    sem = combined.get("semester") or "—"
    dept = combined.get("department_label") or "—"
    total = int(combined.get("total_survey_count") or 0)
    agg = int(combined.get("aggregated_survey_count") or 0)
    paragraphs.append(
        f"يغطي هذا التقرير {total} استبياناً للفصل الدراسي «{sem}» ضمن نطاق «{dept}». "
        f"اكتمل التجميع الإحصائي لـ {agg} استبيان(ات) وفق سياسة الخصوصية والحد الأدنى للإجابات."
    )
    top = combined.get("top_surveys") or []
    bottom = combined.get("bottom_surveys") or []
    if top:
        best = top[0]
        paragraphs.append(
            f"أعلى نتيجة مجمّعة: «{best.get('title_ar')}» بنسبة {best.get('overall_score_percent')}% "
            f"— {interpret_overall_score_ar(best.get('overall_score_percent'), True)}"
        )
    if bottom:
        worst = bottom[0]
        if not top or worst.get("template_code") != top[0].get("template_code"):
            paragraphs.append(
                f"أضعف نتيجة مجمّعة: «{worst.get('title_ar')}» بنسبة {worst.get('overall_score_percent')}% "
                f"— {interpret_overall_score_ar(worst.get('overall_score_percent'), True)}"
            )
    acc_rows = combined.get("accreditation_rows") or []
    gaps = [r for r in acc_rows if (r.get("حالة_الامتثال") or "") in ("فجوة", "جزئي")]
    if gaps:
        indicators = ", ".join(
            {str(r.get("المؤشر") or "") for r in gaps[:5] if r.get("المؤشر")}
        )
        paragraphs.append(
            f"ربط الاعتماد: وُجدت {len(gaps)} حالة امتثال جزئية أو فجوة مرتبطة بمؤشرات "
            f"منها: {indicators or '—'}. يُراجع الجدول التفصيلي أدناه."
        )
    elif acc_rows:
        paragraphs.append(
            "ربط الاعتماد: النتائج المجمّعة تدعم مؤشرات الاعتماد المرتبطة — راجع جدول الربط للتفاصيل."
        )
    ce_secs = combined.get("course_eval_sections") or []
    ce_agg = [s for s in ce_secs if s.get("aggregated")]
    if ce_agg:
        policy = combined.get("course_eval_policy_ar") or format_course_eval_aggregation_policy()
        paragraphs.append(
            f"تقييم المقررات: {len(ce_agg)} شعبة بلغت عتبة التجميع ({policy})."
        )
    return paragraphs


def prepare_combined_pdf_context(
    conn,
    *,
    semester: str | None = None,
    department_id: int | None = None,
    include_course_eval: bool = True,
) -> dict[str, Any]:
    """سياق قالب PDF الموحّد."""
    combined = build_combined_survey_report(
        conn,
        semester=semester,
        department_id=department_id,
        include_course_eval=include_course_eval,
    )
    for r in combined.get("reports") or []:
        r["interpretation_ar"] = interpret_overall_score_ar(
            r.get("overall_score_percent"), r.get("aggregated")
        )
    if combined.get("course_eval"):
        combined["course_eval"] = _enrich_course_eval_report_for_display(combined["course_eval"])
    combined["analysis"] = build_combined_survey_analysis(combined)
    ctx = {
        **combined,
        "executive_summary": _executive_summary_rows(
            combined.get("reports") or [], combined.get("course_eval")
        ),
        "accreditation_rows": _accreditation_map_rows(
            combined.get("reports") or [], combined.get("course_eval")
        ),
        "comparative_rows": _comparative_analysis_rows(
            combined.get("reports") or [], combined.get("course_eval")
        ),
        "metadata_rows": _metadata_rows(combined),
        "narrative_paragraphs": (combined.get("analysis") or {}).get("narrative_paragraphs")
        or generate_executive_narrative_ar(
            {
                **combined,
                "accreditation_rows": _accreditation_map_rows(
                    combined.get("reports") or [], combined.get("course_eval")
                ),
            }
        ),
        "title": "تقرير الاستبيانات الموحّد — ضمان الجودة والاعتماد",
    }
    return enrich_survey_export_context(ctx, for_pdf=False)


def prepare_single_survey_pdf_context(
    conn,
    template_code: str,
    *,
    semester: str | None = None,
    department_id: int | None = None,
) -> dict[str, Any] | None:
    """سياق قالب PDF لاستبيان واحد."""
    code = (template_code or "").strip()
    if code == "student_course":
        report = build_course_eval_report(conn, semester=semester, department_id=department_id)
        report["department_label"] = _department_label(conn, department_id)
        report = _enrich_course_eval_report_for_display(report)
    elif get_template_by_code(conn, code):
        report = build_survey_report(conn, code, semester=semester, department_id=department_id)
        report["interpretation_ar"] = interpret_overall_score_ar(
            report.get("overall_score_percent"), report.get("aggregated")
        )
    else:
        return None
    sem_slug = (report.get("semester") or "report").replace(" ", "_")[:40]
    ctx = {
        "report": report,
        "title": f"تقرير {report.get('title_ar') or code}",
        "metadata_rows": [
            {"البند": "الفصل الدراسي", "القيمة": report.get("semester")},
            {"البند": "النطاق", "القيمة": report.get("department_label")},
            {"البند": "رمز الاستبيان", "القيمة": report.get("template_code")},
            {"البند": "فئة المستجيب", "القيمة": report.get("respondent_label")},
            {"البند": "عدد الإجابات", "القيمة": report.get("response_count")},
            {"البند": "الحد الأدنى للتجميع", "القيمة": report.get("min_aggregate")},
            {
                "البند": "حالة التجميع",
                "القيمة": "مكتمل" if report.get("aggregated") else "ناقص",
            },
            {"البند": "النتيجة الإجمالية %", "القيمة": report.get("overall_score_percent")},
            {"البند": "حالة الامتثال", "القيمة": report.get("compliance_status_ar")},
            {
                "البند": "الخصوصية",
                "القيمة": "لا تُعرض إجابات فردية — التجميع بعد الحد الأدنى فقط.",
            },
        ],
        "filename_prefix": f"survey_{code}_{sem_slug}",
    }
    return enrich_survey_export_context(ctx, for_pdf=False)


def prepare_course_eval_section_pdf_context(
    conn,
    section_id: int,
    *,
    semester: str | None = None,
    department_id: int | None = None,
) -> dict[str, Any] | None:
    """سياق معاينة/PDF لتقييم شعبة واحدة."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    report = build_course_eval_section_report(
        conn,
        int(section_id),
        semester=sem,
        department_id=department_id,
    )
    if not report:
        return None
    report.setdefault(
        "department_label",
        report.get("department_name") or _department_label(conn, department_id),
    )
    report = _enrich_course_eval_report_for_display(report)
    sid = int(report.get("section_id") or section_id)
    sem_slug = sem.replace(" ", "_")[:40]
    scope_bits = [
        f"المقرر: {report.get('course_name') or '—'}",
        f"الشعبة: {sid}",
        f"الأستاذ: {report.get('instructor_name') or '—'}",
    ]
    if report.get("group_code_label"):
        scope_bits.append(f"المجموعة: {report.get('group_code_label')}")
    report["scope_note_ar"] = " — ".join(scope_bits)
    ctx = {
        "report": report,
        "title": f"تقرير {report.get('title_ar') or 'تقييم المقرر والأستاذ'}",
        "metadata_rows": [
            {"البند": "الفصل الدراسي", "القيمة": report.get("semester")},
            {"البند": "المقرر", "القيمة": report.get("course_name")},
            {"البند": "الشعبة", "القيمة": sid},
            {"البند": "الأستاذ", "القيمة": report.get("instructor_name")},
            {"البند": "القسم", "القيمة": report.get("department_name") or report.get("department_label")},
            {"البند": "مسجّلون (تقدير)", "القيمة": report.get("enrolled_count")},
            {"البند": "عدد التقييمات", "القيمة": report.get("response_count")},
            {"البند": "نسبة الاستجابة %", "القيمة": report.get("response_rate_percent")},
            {"البند": "الحد الأدنى للتجميع", "القيمة": report.get("min_aggregate")},
            {
                "البند": "حالة التجميع",
                "القيمة": "مكتمل" if report.get("aggregated") else "ناقص — لا تُعرض تفاصيل البنود",
            },
            {"البند": "النتيجة الإجمالية %", "القيمة": report.get("overall_score_percent")},
            {"البند": "حالة الامتثال", "القيمة": report.get("compliance_status_ar")},
            {
                "البند": "الخصوصية",
                "القيمة": "لا تُعرض إجابات فردية — التجميع بعد بلوغ نسبة المسجّلين فقط.",
            },
        ],
        "filename_prefix": f"course_eval_section_{sid}_{sem_slug}",
        "preview_banner_title": "معاينة تقييم الشعبة",
    }
    return enrich_survey_export_context(ctx, for_pdf=False)


def prepare_course_eval_by_course_pdf_context(
    conn,
    course_name: str,
    instructor_id: int,
    *,
    semester: str | None = None,
    department_id: int | None = None,
) -> dict[str, Any] | None:
    """سياق معاينة/PDF لتجميع مقرر + أستاذ عبر شعب متعددة."""
    sem = (semester or "").strip() or term_label_from_conn(conn)
    cname = (course_name or "").strip()
    iid = int(instructor_id or 0)
    if not cname or iid <= 0:
        return None
    report = build_course_eval_by_course_report(
        conn,
        cname,
        iid,
        semester=sem,
        department_id=department_id,
    )
    if not report:
        return None
    report.setdefault(
        "department_label",
        report.get("department_name") or _department_label(conn, department_id),
    )
    report = _enrich_course_eval_report_for_display(report)
    section_ids = report.get("section_ids") or []
    sections_txt = "، ".join(str(x) for x in section_ids) if section_ids else "—"
    sem_slug = sem.replace(" ", "_")[:40]
    safe_course = cname.replace(" ", "_")[:30]
    report["scope_note_ar"] = (
        f"المقرر: {cname} — الأستاذ: {report.get('instructor_name') or '—'} — الشعب: {sections_txt}"
    )
    ctx = {
        "report": report,
        "title": f"تقرير {report.get('title_ar') or 'تقييم المقرر والأستاذ'}",
        "metadata_rows": [
            {"البند": "الفصل الدراسي", "القيمة": report.get("semester")},
            {"البند": "المقرر", "القيمة": report.get("course_name")},
            {"البند": "الأستاذ", "القيمة": report.get("instructor_name")},
            {"البند": "عدد الشعب", "القيمة": report.get("section_count")},
            {"البند": "أرقام الشعب", "القيمة": sections_txt},
            {"البند": "القسم", "القيمة": report.get("department_name") or report.get("department_label")},
            {"البند": "مسجّلون (تقدير)", "القيمة": report.get("enrolled_count")},
            {"البند": "عدد التقييمات", "القيمة": report.get("response_count")},
            {"البند": "نسبة الاستجابة %", "القيمة": report.get("response_rate_percent")},
            {"البند": "الحد الأدنى للتجميع", "القيمة": report.get("min_aggregate")},
            {
                "البند": "حالة التجميع",
                "القيمة": "مكتمل" if report.get("aggregated") else "ناقص — لا تُعرض تفاصيل البنود",
            },
            {"البند": "النتيجة الإجمالية %", "القيمة": report.get("overall_score_percent")},
            {"البند": "حالة الامتثال", "القيمة": report.get("compliance_status_ar")},
            {
                "البند": "الخصوصية",
                "القيمة": "لا تُعرض إجابات فردية — التجميع بعد بلوغ نسبة المسجّلين فقط.",
            },
        ],
        "filename_prefix": f"course_eval_{safe_course}_inst{iid}_{sem_slug}",
        "preview_banner_title": "معاينة تقييم المقرر + الأستاذ",
    }
    return enrich_survey_export_context(ctx, for_pdf=False)
