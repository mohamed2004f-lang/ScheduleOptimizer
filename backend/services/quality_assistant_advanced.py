"""ميزات متقدمة لمساعد الجودة: ملخص لجنة، أرشيف، مدقق صياغة، تنبيهات، استخدام، LLM اختياري، حزمة أسلوب."""

from __future__ import annotations

import io
import json
import os
import re
import zipfile
from typing import Any
from urllib import error as urlerror
from urllib import request as urlrequest

from backend.core.department_archive_catalog import (
    ARCHIVE_RECORD_TYPES,
    suggestions_for_type,
)
from backend.core.quality_assistant_catalog import POLICY_BANNER_AR
from backend.database.database import is_postgresql
from backend.services.quality_metrics import term_label_from_conn


def _now() -> str:
    from datetime import datetime, timezone

    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def ensure_usage_tables(conn) -> None:
    cur = conn.cursor()
    if is_postgresql():
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS quality_assistant_usage_events (
                id BIGSERIAL PRIMARY KEY,
                mode TEXT NOT NULL DEFAULT '',
                intent TEXT NOT NULL DEFAULT '',
                channel TEXT NOT NULL DEFAULT '',
                page_path TEXT NOT NULL DEFAULT '',
                department_id BIGINT,
                actor_hash TEXT NOT NULL DEFAULT '',
                meta_json TEXT NOT NULL DEFAULT '{}',
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
    else:
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS quality_assistant_usage_events (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                mode TEXT NOT NULL DEFAULT '',
                intent TEXT NOT NULL DEFAULT '',
                channel TEXT NOT NULL DEFAULT '',
                page_path TEXT NOT NULL DEFAULT '',
                department_id INTEGER,
                actor_hash TEXT NOT NULL DEFAULT '',
                meta_json TEXT NOT NULL DEFAULT '{}',
                created_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
    conn.commit()


def _actor_hash(actor: str) -> str:
    import hashlib

    raw = (actor or "").strip().encode("utf-8")
    if not raw:
        return "anon"
    return hashlib.sha256(raw).hexdigest()[:16]


def log_usage_event(
    conn,
    *,
    mode: str = "",
    intent: str = "",
    channel: str = "",
    page_path: str = "",
    department_id: int | None = None,
    actor: str = "",
    meta: dict[str, Any] | None = None,
) -> None:
    ensure_usage_tables(conn)
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO quality_assistant_usage_events
        (mode, intent, channel, page_path, department_id, actor_hash, meta_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            (mode or "")[:64],
            (intent or "")[:64],
            (channel or "")[:32],
            (page_path or "")[:260],
            int(department_id) if department_id is not None else None,
            _actor_hash(actor),
            json.dumps(meta or {}, ensure_ascii=False)[:2000],
            _now(),
        ),
    )
    conn.commit()


def usage_analytics_summary(conn, *, limit: int = 20) -> dict[str, Any]:
    ensure_usage_tables(conn)
    cur = conn.cursor()
    rows = cur.execute(
        """
        SELECT intent, mode, channel, COUNT(*) AS c
        FROM quality_assistant_usage_events
        GROUP BY intent, mode, channel
        ORDER BY c DESC
        LIMIT ?
        """,
        (int(limit),),
    ).fetchall() or []
    top = []
    for r in rows:
        if hasattr(r, "keys"):
            top.append(
                {
                    "intent": r["intent"],
                    "mode": r["mode"],
                    "channel": r["channel"],
                    "count": int(r["c"]),
                }
            )
        else:
            top.append(
                {
                    "intent": r[0],
                    "mode": r[1],
                    "channel": r[2],
                    "count": int(r[3]),
                }
            )
    total_row = cur.execute("SELECT COUNT(*) FROM quality_assistant_usage_events").fetchone()
    total = int(total_row[0] if not hasattr(total_row, "keys") else list(total_row)[0])
    return {
        "status": "ok",
        "total_events": total,
        "top": top,
        "suggestion_only": True,
        "note_ar": "سجل مجهول نسبياً (hash للمستخدم) لتحسين المواضيع والدليل — ليس تقييماً للأفراد.",
    }


def llm_config() -> dict[str, Any]:
    enabled = (os.environ.get("QUALITY_ASSISTANT_LLM_ENABLED") or "").strip().lower() in (
        "1",
        "true",
        "yes",
        "on",
    )
    return {
        "enabled": enabled,
        "base_url": (os.environ.get("QUALITY_ASSISTANT_LLM_BASE_URL") or "").rstrip("/"),
        "api_key": (os.environ.get("QUALITY_ASSISTANT_LLM_API_KEY") or "").strip(),
        "model": (os.environ.get("QUALITY_ASSISTANT_LLM_MODEL") or "gpt-4o-mini").strip(),
        "timeout_sec": int(os.environ.get("QUALITY_ASSISTANT_LLM_TIMEOUT") or "25"),
    }


def optional_llm_complete(
    *,
    system_ar: str,
    user_ar: str,
    temperature: float = 0.3,
) -> dict[str, Any]:
    """استدعاء موصل OpenAI-compatible اختياري. لا يُستخدم إن لم يُفعَّل."""
    cfg = llm_config()
    if not cfg["enabled"] or not cfg["base_url"] or not cfg["api_key"]:
        return {"ok": False, "skipped": True, "text": "", "message_ar": "موصل LLM غير مفعّل."}
    payload = {
        "model": cfg["model"],
        "temperature": temperature,
        "messages": [
            {"role": "system", "content": system_ar},
            {"role": "user", "content": user_ar},
        ],
    }
    req = urlrequest.Request(
        cfg["base_url"] + "/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {cfg['api_key']}",
        },
        method="POST",
    )
    try:
        with urlrequest.urlopen(req, timeout=cfg["timeout_sec"]) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        text = (
            (((data.get("choices") or [{}])[0].get("message") or {}).get("content")) or ""
        ).strip()
        return {"ok": bool(text), "skipped": False, "text": text, "provider": "openai_compatible"}
    except (urlerror.URLError, urlerror.HTTPError, TimeoutError, json.JSONDecodeError, OSError) as ex:
        return {"ok": False, "skipped": False, "text": "", "error": str(ex)[:300]}


def build_committee_summary(
    conn,
    *,
    mode: str,
    semester: str | None = None,
    department_id: int | None = None,
    notes: str = "",
    history: list[dict[str, str]] | None = None,
) -> dict[str, Any]:
    from backend.services.quality_assistant import build_context, normalize_chat_history

    hist = normalize_chat_history(history or [])
    ctx = build_context(conn, mode=mode, semester=semester, department_id=department_id)
    sem = ctx.get("semester") or term_label_from_conn(conn)
    dept = ctx.get("department") or {}
    agenda: list[str] = []
    for g in ((ctx.get("prog") or {}).get("top_gaps") or [])[:5]:
        agenda.append(f"PROG {g.get('code')}: {g.get('title_ar')} [{g.get('status')}]")
    for g in ((ctx.get("inst") or {}).get("top_gaps") or [])[:5]:
        agenda.append(f"INST {g.get('code')}: {g.get('title_ar')} [{g.get('status')}]")
    for tip in ((ctx.get("archive") or {}).get("tips_ar") or [])[:4]:
        if "لا توجد فجوات" not in tip:
            agenda.append(f"أرشيف: {tip}")
    discussion = []
    for h in hist[-8:]:
        who = "عضو" if h.get("role") == "user" else "مساعد"
        discussion.append(f"- {who}: {h.get('text')}")
    if notes.strip():
        discussion.append(f"- ملاحظات الجلسة: {notes.strip()}")

    recommendations = [
        "مراجعة الفجوات أعلى أولوية وربط الشواهد يدوياً.",
        "تأكيد بشري لأي توصية قبل الاعتماد أو الإغلاق.",
        "توثيق المتابعة في أرشيف القسم (محضر/قرار/ملاحظة).",
    ]
    if not agenda:
        agenda = ["لا فجوات بارزة من السياق — راجع لوحة الجودة يدوياً."]

    sections = {
        "title_ar": "مسودة ملخص لجنة الجودة العلمية",
        "semester": sem,
        "department_name_ar": dept.get("name_ar") or "الكلية",
        "department_code": dept.get("code") or "",
        "agenda": agenda,
        "discussion": discussion or ["- (لا محادثة مسجّلة — أضف ملاحظات الجلسة)"],
        "recommendations": recommendations,
        "disclaimer_ar": (
            "مسودة غير ملزمة — للاستخدام الداخلي قبل الاعتماد البشري. " + POLICY_BANNER_AR
        ),
    }
    md_lines = [
        f"# {sections['title_ar']}",
        "",
        f"- الفصل: {sem}",
        f"- القسم/النطاق: {sections['department_name_ar']} ({sections['department_code'] or '—'})",
        "",
        "> " + sections["disclaimer_ar"],
        "",
        "## بنود الأجندة المقترحة",
        *[f"1. {a}" if False else f"- {a}" for a in agenda],
        "",
        "## أبرز النقاش",
        *discussion,
        "",
        "## توصيات للمتابعة (مسودة)",
        *[f"- {r}" for r in recommendations],
        "",
    ]
    markdown = "\n".join(md_lines)
    return {
        "status": "ok",
        "sections": sections,
        "markdown": markdown,
        "suggestion_only": True,
        "policy_ar": POLICY_BANNER_AR,
    }


def committee_summary_docx_bytes(summary: dict[str, Any]) -> bytes:
    try:
        from docx import Document
    except ImportError as ex:
        raise RuntimeError("python-docx غير متاح") from ex
    sec = summary.get("sections") or {}
    doc = Document()
    doc.add_heading(sec.get("title_ar") or "ملخص لجنة", level=1)
    doc.add_paragraph(f"الفصل: {sec.get('semester') or '—'}")
    doc.add_paragraph(
        f"النطاق: {sec.get('department_name_ar') or '—'} ({sec.get('department_code') or '—'})"
    )
    doc.add_paragraph(sec.get("disclaimer_ar") or POLICY_BANNER_AR)
    doc.add_heading("بنود الأجندة المقترحة", level=2)
    for a in sec.get("agenda") or []:
        doc.add_paragraph(str(a), style="List Bullet")
    doc.add_heading("أبرز النقاش", level=2)
    for d in sec.get("discussion") or []:
        doc.add_paragraph(str(d))
    doc.add_heading("توصيات للمتابعة (مسودة)", level=2)
    for r in sec.get("recommendations") or []:
        doc.add_paragraph(str(r), style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def committee_summary_pdf_bytes(summary: dict[str, Any]) -> bytes | None:
    from backend.core.chromium_pdf import pdf_bytes_from_html_chromium

    sec = summary.get("sections") or {}
    agenda_html = "".join(f"<li>{_esc(a)}</li>" for a in (sec.get("agenda") or []))
    disc_html = "".join(f"<li>{_esc(d)}</li>" for d in (sec.get("discussion") or []))
    rec_html = "".join(f"<li>{_esc(r)}</li>" for r in (sec.get("recommendations") or []))
    html = f"""<!doctype html><html lang="ar" dir="rtl"><head><meta charset="utf-8">
    <style>
      body {{ font-family: 'Segoe UI', Tahoma, Arial, sans-serif; margin: 2rem; line-height: 1.6; }}
      h1 {{ font-size: 1.4rem; }} h2 {{ font-size: 1.1rem; margin-top: 1.2rem; }}
      .warn {{ background: #fff8e6; border: 1px solid #f0d78c; padding: .75rem; }}
    </style></head><body>
    <h1>{_esc(sec.get('title_ar') or 'ملخص لجنة')}</h1>
    <p>الفصل: {_esc(sec.get('semester'))} · النطاق: {_esc(sec.get('department_name_ar'))}</p>
    <div class="warn">{_esc(sec.get('disclaimer_ar'))}</div>
    <h2>بنود الأجندة المقترحة</h2><ul>{agenda_html}</ul>
    <h2>أبرز النقاش</h2><ul>{disc_html}</ul>
    <h2>توصيات للمتابعة (مسودة)</h2><ul>{rec_html}</ul>
    </body></html>"""
    data, err = pdf_bytes_from_html_chromium(html)
    if err or not data:
        return None
    return data


def _esc(s: Any) -> str:
    t = str(s or "")
    return (
        t.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def suggest_archive_links(
    conn,
    *,
    mode: str,
    semester: str | None = None,
    department_id: int | None = None,
    notes: str = "",
) -> dict[str, Any]:
    """اقتراح نوع سجل أرشيف + مؤشرات QAA مرشّحة للفجوات — ربط يدوي فقط."""
    from backend.services.quality_assistant import build_context

    ctx = build_context(conn, mode=mode, semester=semester, department_id=department_id)
    suggestions: list[dict[str, Any]] = []
    q = (notes or "").strip()

    def add(gap_label: str, record_type: str, why_ar: str):
        meta = ARCHIVE_RECORD_TYPES.get(record_type) or {}
        suggestions.append(
            {
                "gap_label_ar": gap_label,
                "record_type": record_type,
                "record_type_label_ar": meta.get("title_ar") or record_type,
                "why_ar": why_ar,
                "qaa_suggestions": suggestions_for_type(record_type)[:4],
                "href": "/academic_quality/archive",
                "suggestion_only": True,
            }
        )

    for g in ((ctx.get("prog") or {}).get("top_gaps") or [])[:4]:
        code = str(g.get("code") or "")
        title = str(g.get("title_ar") or "")
        rtype = "minutes"
        blob = (code + " " + title).lower()
        if any(k in blob for k in ("قرار", "اعتماد", "سياسة")):
            rtype = "decision"
        elif any(k in blob for k in ("مراسل", "خطاب", "كتاب")):
            rtype = "corr_out"
        elif any(k in blob for k in ("تدريب", "ملاحظة", "متابعة")):
            rtype = "notes"
        add(
            f"PROG {code}: {title}",
            rtype,
            "فجوة برامجية ظاهرة — اقترح توثيقها بمحضر/قرار ثم ربط الشاهد يدوياً.",
        )

    for g in ((ctx.get("inst") or {}).get("top_gaps") or [])[:3]:
        add(
            f"INST {g.get('code')}: {g.get('title_ar')}",
            "minutes",
            "فجوة مؤسسية — غالباً تُغطى بمحضر لجنة + شواهد مرفقة بعد التأكيد البشري.",
        )

    for tip in ((ctx.get("archive") or {}).get("tips_ar") or [])[:3]:
        if "لا توجد فجوات" in tip:
            continue
        rtype = "notes"
        if "محضر" in tip:
            rtype = "minutes"
        elif "قرار" in tip:
            rtype = "decision"
        add(tip, rtype, "من نواقص الأرشيف الحالية — أكمل السجل ثم اربطه يدوياً إن لزم.")

    if q:
        ql = q.lower()
        if any(k in ql for k in ("محضر", "اجتماع", "لجنة")):
            add(q[:180], "minutes", "من نص سؤالك — يبدو مناسباً لمحضر.")
        elif any(k in ql for k in ("قرار", "اعتماد")):
            add(q[:180], "decision", "من نص سؤالك — يبدو مناسباً لقرار.")
        elif any(k in ql for k in ("صادر", "خطاب", "مراسل")):
            add(q[:180], "corr_out", "من نص سؤالك — يبدو مناسباً لمراسلة.")

    if not suggestions:
        add(
            "لا فجوة محددة",
            "minutes",
            "ابدأ بمحضر جلسة جودة دوري كحد أدنى تشغيلي.",
        )

    return {
        "status": "ok",
        "suggestions": suggestions[:12],
        "message_ar": "اقتراحات ربط أرشيف (يدوي فقط — لا ربط تلقائي).",
        "suggestion_only": True,
        "policy_ar": POLICY_BANNER_AR,
        "links": [{"href": "/academic_quality/archive", "label_ar": "أرشيف القسم"}],
    }


def proofread_quality_text(
    *,
    text: str,
    kind: str = "auto",
    use_llm: bool = True,
) -> dict[str, Any]:
    """مدقق صياغة لرسالة/رؤية/CLO — قواعد محلية + LLM اختياري."""
    raw = (text or "").strip()
    if not raw:
        return {
            "status": "ok",
            "message_ar": "الصق الفقرة المراد تدقيقها في مربع الملاحظات.",
            "issues": [],
            "improved_ar": "",
            "suggestion_only": True,
        }
    kind_l = (kind or "auto").strip().lower()
    if kind_l == "auto":
        if any(k in raw for k in ("أن يكون", "CLO", "clo", "يتمكن", "يصمم", "يحلّل", "يطبّق")):
            kind_l = "clo"
        elif any(k in raw for k in ("رؤية", "نطمح", "بحلول عام")):
            kind_l = "vision"
        else:
            kind_l = "mission"

    issues: list[str] = []
    improved = raw
    if len(raw) < 25:
        issues.append("النص قصير جداً — أضف الغرض والجمهور والنتيجة المتوقعة.")
    if re.search(r"(ممتاز|الأفضل|عالمي|رائد|متميز)\b", raw) and not re.search(
        r"\d{4}|مؤشر|نسبة|%|خلال", raw
    ):
        issues.append("صياغة فضفاضة دون مؤشر زمني أو قابل للقياس.")
    vague = ("جودة عالية", "بشكل ممتاز", "أفضل الممارسات", "على أعلى مستوى")
    for v in vague:
        if v in raw:
            issues.append(f"تجنّب العبارة العامة: «{v}» — استبدلها بنتيجة قابلة للتحقق.")

    if kind_l == "clo":
        if not re.search(
            r"^(يصمم|يحلل|يحلّل|يطبق|يطبّق|يقيم|يقيّم|ينشئ|يشرح|يحدد|يستخدم|يركب)",
            raw.strip(),
        ):
            issues.append("يفضّل أن يبدأ CLO بفعل قابل للملاحظة (يصمم، يحلّل، يطبّق…).")
            if not improved.lower().startswith(("يصمم", "يحلل", "يحلّل", "يطبق", "يطبّق")):
                improved = "يصمم / يحلّل / يطبّق — " + raw
        if "الطالب" not in raw and "المتعلم" not in raw:
            issues.append("وضّح فاعل المخرج (الطالب/المتعلم) إن أمكن.")
    elif kind_l == "mission":
        if not any(k in raw for k in ("إعداد", "تأهيل", "خدمة", "تطوير", "تقديم")):
            issues.append("الرسالة عادة تصف الغرض الحالي والجمهور والنتيجة.")
    elif kind_l == "vision":
        if not re.search(r"20\d{2}|خلال|بحلول", raw):
            issues.append("الرؤية أوضح بوجود أفق زمني.")

    issues.append("راجع محلياً مقابل مؤشرات QAA-2023.4-INST / PROG-UG قبل الاعتماد.")
    llm_note = ""
    if use_llm:
        llm = optional_llm_complete(
            system_ar=(
                "أنت مدقق صياغة أكاديمية للكلية. حسّن النص دون ادعاء امتثال. "
                "أجب بالعربية فقط بنص محسّن مختصر ثم سطر Issues: قائمة نقاط."
            ),
            user_ar=f"النوع: {kind_l}\nالنص:\n{raw}",
        )
        if llm.get("ok") and llm.get("text"):
            improved = llm["text"]
            llm_note = "شُغّل موصل LLM الاختياري فوق القواعد المحلية."
        elif llm.get("enabled") is False or llm.get("skipped"):
            llm_note = "القواعد المحلية فقط (LLM غير مفعّل)."
        else:
            llm_note = "تعذّر LLM — أعيدت القواعد المحلية."

    return {
        "status": "ok",
        "kind": kind_l,
        "issues": issues,
        "improved_ar": improved,
        "original_ar": raw,
        "llm_note_ar": llm_note,
        "message_ar": f"مدقق صياغة ({kind_l}) — اقتراح فقط.",
        "suggestion_only": True,
        "policy_ar": POLICY_BANNER_AR,
        "disclaimer_ar": "التحسين مساعدة صياغة ولا يعتمد امتثالاً ذاتياً.",
    }


def proactive_term_alerts(
    conn,
    *,
    mode: str,
    semester: str | None = None,
    department_id: int | None = None,
) -> dict[str, Any]:
    from backend.services.quality_assistant import build_welcome_brief

    welcome = build_welcome_brief(
        conn, mode=mode, semester=semester, department_id=department_id
    )
    alerts = []
    for t in welcome.get("tasks") or []:
        alerts.append(
            {
                "severity_ar": "تحذير",
                "title_ar": t.get("title_ar") or "",
                "href": t.get("href") or "/academic_quality/dashboard",
                "source": t.get("source") or "context",
            }
        )
    # تنبيه زمني عام قبل الإغلاق
    alerts.append(
        {
            "severity_ar": "تذكير",
            "title_ar": "قبل إغلاق الفصل: راجع إغلاق المقررات، الأرشيف، وربط شواهد PROG/INST يدوياً.",
            "href": "/academic_quality/dashboard",
            "source": "term_close",
        }
    )
    # إزالة تكرار
    uniq = []
    seen = set()
    for a in alerts:
        k = (a.get("title_ar") or "")[:100]
        if k in seen:
            continue
        seen.add(k)
        uniq.append(a)
    return {
        "status": "ok",
        "alerts": uniq[:8],
        "message_ar": "تنبيهات استباقية قبل إغلاق الفصل (اقتراح متابعة).",
        "semester": welcome.get("semester"),
        "suggestion_only": True,
        "policy_ar": POLICY_BANNER_AR,
    }


def build_style_training_export(conn) -> bytes:
    """
    حزمة أسلوب/بيانات مرخّصة للتدريب المستقبلي (ليس تشغيل LoRA هنا).
    تتضمن ملخصات مكتبة المعرفة المعتمدة وإرشادات الأسلوب.
    """
    from backend.services.quality_knowledge import (
        ensure_quality_knowledge_tables,
        get_knowledge_doc,
        list_knowledge_docs,
    )

    ensure_quality_knowledge_tables(conn)
    docs = list_knowledge_docs(conn, status="approved", limit=300)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(
            "README_AR.md",
            "\n".join(
                [
                    "# حزمة أسلوب الكلية — للتدريب المستقبلي الاختياري",
                    "",
                    "هذه الحزمة تجمع وثائق **معتمدة** تملك الكلية حق استخدامها.",
                    "ليست تشغيلاً لتدريب LoRA داخل المنظومة.",
                    "يمكن لاحقاً استخدامها خارجياً لـ adapter خاص بالكلية وفق سياسة الخصوصية والترخيص.",
                    "",
                    POLICY_BANNER_AR,
                ]
            ),
        )
        zf.writestr(
            "style_prompt_ar.txt",
            "\n".join(
                [
                    "أسلوب مساعد الجودة:",
                    "- اقتراح فقط دون اعتماد امتثال.",
                    "- المرجع الملزم: QAA ليبيا INST/PROG.",
                    "- المراجع العالمية للمناقشة والصياغة فقط.",
                    "- اربط الشواهد يدوياً بعد التأكيد البشري.",
                    "- لغة عربية رسمية مختصرة وقابلة للتنفيذ.",
                ]
            ),
        )
        manifest = []
        for d in docs:
            full = get_knowledge_doc(conn, int(d["id"])) or d
            text = (full.get("extracted_text") or "")[:50000]
            safe = re.sub(r"[^\w\-]+", "_", f"{d['id']}_{d.get('title_ar') or 'doc'}")[:80]
            zf.writestr(f"corpus/{safe}.md", f"# {d.get('title_ar')}\n\n{text}\n")
            manifest.append(
                {
                    "id": d.get("id"),
                    "title_ar": d.get("title_ar"),
                    "category": d.get("category"),
                    "source_label_ar": d.get("source_label_ar"),
                }
            )
        zf.writestr(
            "manifest.json",
            json.dumps(
                {
                    "docs_count": len(manifest),
                    "docs": manifest,
                    "purpose": "future_optional_finetune_corpus",
                    "not_in_app_training": True,
                },
                ensure_ascii=False,
                indent=2,
            ),
        )
        # إحصاء استخدام إن وُجد
        try:
            usage = usage_analytics_summary(conn, limit=50)
            zf.writestr(
                "usage_top.json", json.dumps(usage, ensure_ascii=False, indent=2)
            )
        except Exception:
            pass
    return buf.getvalue()


def enhance_with_optional_llm(base_reply: dict[str, Any], *, user_question: str) -> dict[str, Any]:
    """إن فُعّل LLM: يضيف فقرة صياغة أغنى دون تغيير سياسة الاقتراح."""
    if not (user_question or "").strip():
        return base_reply
    bullets = base_reply.get("bullets") or []
    context_snip = "\n".join(str(b) for b in bullets[:12])
    llm = optional_llm_complete(
        system_ar=(
            "أنت مساعد جودة أكاديمية. أعد صياغة إجابة عربية قصيرة قابلة للتنفيذ. "
            "لا تدّعِ اعتماداً أو امتثالاً. اذكر أن المرجع الملزم QAA المحلي."
        ),
        user_ar=f"سؤال المستخدم:\n{user_question}\n\nمسودة النظام:\n{context_snip}",
    )
    out = dict(base_reply)
    if llm.get("ok") and llm.get("text"):
        out["llm_enrichment_ar"] = llm["text"]
        out["knowledge_tag"] = (out.get("knowledge_tag") or "") + " · LLM اختياري"
        bullets2 = list(bullets)
        bullets2.extend(["", "صياغة مساندة (LLM اختياري — غير ملزمة):", llm["text"]])
        out["bullets"] = bullets2
    else:
        out["llm_enrichment_ar"] = ""
        out["llm_status"] = "skipped" if llm.get("skipped") else "error"
    return out
