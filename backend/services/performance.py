from datetime import datetime

from flask import Blueprint, jsonify, render_template, request, send_file, session

from backend.core.auth import role_required
from .utilities import get_connection, excel_response_from_df, pdf_response_from_html, get_current_term
from backend.services.grades import _load_transcript_data


performance_bp = Blueprint("performance", __name__)


def _load_rule_number(conn, rule_key: str, default: float) -> float:
    """
    قراءة قيمة رقمية من جدول academic_rules إن وُجدت، وإلا إرجاع القيمة الافتراضية.
    """
    try:
        cur = conn.cursor()
        row = cur.execute(
            "SELECT value_number FROM academic_rules WHERE rule_key = ? AND is_active = 1",
            (rule_key,),
        ).fetchone()
        if not row:
            return float(default)
        val = row[0]
        if val is None:
            return float(default)
        return float(val)
    except Exception:
        return float(default)


def _compute_status(ordered_semesters, semester_gpas, cumulative_gpa):
    """
    حساب حالة الطالب (جيد، إنذار أول، إنذار ثانٍ، أكثر) بشكل تقريبي بناءً على المادة 43.
    - نتجاهل أول فصل قُيّد فيه الطالب (لا يصدر فيه إنذار).
    - نعتبر الفصل "منخفضاً" إذا كان معدله الفصلي أقل من الحد المحدد في لائحة الإنذار (افتراضياً 50%).
    - نحسب عدد الفصول المنخفضة المتتالية في النهاية:
        0 -> good
        1 -> warning_1
        2 -> warning_2
        3+ -> warning_3
    - إذا كان المعدل التراكمي أقل من حد الفصل (افتراضياً 35%) بعد مرور فصلين أو أكثر، نضيف ملاحظة عن احتمال الفصل.
    """
    if not ordered_semesters:
        return {"code": "no_data", "label": "لا توجد بيانات درجات"}

    # قراءة الحدود الرقمية من جدول academic_rules (مع قيم افتراضية إذا لم توجد)
    with get_connection() as conn:
        warning_threshold = _load_rule_number(conn, "warning_semester_threshold", 50.0)
        dismissal_cgpa_threshold = _load_rule_number(conn, "dismissal_cgpa_threshold", 35.0)
        dismissal_min_semesters = int(_load_rule_number(conn, "dismissal_min_semesters", 2.0))
        # حدود مدة الدراسة
        study_normal_max = int(_load_rule_number(conn, "study_normal_max_semesters", 10.0))
        study_abs_max = int(_load_rule_number(conn, "study_absolute_max_semesters", 14.0))
        study_extra_semesters = int(_load_rule_number(conn, "study_extra_semesters_once", 2.0))
        study_extra_min_units = int(_load_rule_number(conn, "study_extra_semesters_min_units", 130.0))

    lows = []
    for idx, sem in enumerate(ordered_semesters):
        g = semester_gpas.get(sem, 0.0)
        if idx == 0:
            lows.append(False)
        else:
            lows.append(g < warning_threshold)

    consecutive_lows = 0
    for idx in range(len(lows) - 1, -1, -1):
        if not lows[idx]:
            break
        if idx == 0:
            break
        consecutive_lows += 1

    if consecutive_lows == 0:
        status_code = "good"
        label = "طالب في وضع أكاديمي سليم"
    elif consecutive_lows == 1:
        status_code = "warning_1"
        label = f"إنذار أكاديمي أول (معدل فصلي أقل من {warning_threshold:.0f}%)"
    elif consecutive_lows == 2:
        status_code = "warning_2"
        label = "إنذار أكاديمي ثانٍ (فصلان متتاليان دون إزالة الإنذار)"
    else:
        status_code = "warning_3"
        label = "أكثر من إنذارين متتاليين (يستدعي دراسة حالة للفصل المحتمل)"

    extra_notes = []
    try:
        cgpa = float(cumulative_gpa or 0.0)
    except Exception:
        cgpa = 0.0

    semesters_count = len(ordered_semesters)
    if semesters_count and cgpa < dismissal_cgpa_threshold:
        if semesters_count < dismissal_min_semesters:
            # مرحلة مبكرة: تنبيه فقط دون حديث مباشر عن الفصل
            extra_notes.append(
                f"المعدل التراكمي أقل من {dismissal_cgpa_threshold:.0f}% في هذه المرحلة المبكرة من الدراسة؛ "
                f"يُنصح الطالب بتحسين أدائه لتفادي الوصول إلى حد الفصل وفق اللائحة."
            )
        else:
            # بعد تجاوز عدد الفصول المحدد: إشارة صريحة لاحتمال الفصل
            extra_notes.append(
                f"المعدل التراكمي أقل من {dismissal_cgpa_threshold:.0f}% بعد {semesters_count} فصل/فصول دراسية منذ الالتحاق؛ "
                f"وفق المادة 40 أو ما يعادلها قد يُعرض الطالب للفصل، مع إمكانية منحه فرصة استثنائية واحدة حسب اللوائح."
            )

    # تنبيهات خاصة بمدة الدراسة (عدد الفصول منذ الالتحاق)
    if semesters_count:
        if semesters_count >= study_abs_max:
            extra_notes.append(
                f"تجاوز الطالب الحد الأقصى المطلق لمدة الدراسة ({study_abs_max} فصل دراسي اعتيادي أو أكثر)؛ "
                f"يجب عرض حالته على مجلس الكلية لاتخاذ قرار نهائي."
            )
        elif semesters_count >= max(study_abs_max - 2, study_normal_max):
            remaining = study_abs_max - semesters_count
            if remaining > 0:
                extra_notes.append(
                    f"تنبيه مدة الدراسة: تبقّى تقريباً {remaining} فصل/فصول قبل بلوغ الحد الأقصى المطلق لمدة الدراسة ({study_abs_max} فصلاً)."
                )
        elif semesters_count >= study_normal_max:
            extra_notes.append(
                f"تنبيه مدة الدراسة: عدد الفصول المنجزة ({semesters_count}) تجاوز المدة النظامية ({study_normal_max} فصول)، "
                f"لكن لم يصل بعد إلى الحد الأقصى المطلق ({study_abs_max}). يمكن النظر في منح فصول إضافية (حتى {study_extra_semesters} فصل/فصول) "
                f"لطالب اجتاز {study_extra_min_units} وحدة دراسية على الأقل وبتوصية من القسم."
            )

    if extra_notes:
        label = f"{label} — " + " ".join(extra_notes)

    return {"code": status_code, "label": label}


def _override_status_by_enrollment(enrollment_status: str, academic_status: dict) -> dict:
    """
    دمج حالة القيد (سحب ملف/إيقاف قيد) مع حالة الإنذار الأكاديمي.
    إذا كانت حالة القيد withdrawn/suspended نعرضها في تقرير الأداء بدل الأكاديمي.
    """
    try:
        es = (enrollment_status or "").strip().lower()
    except Exception:
        es = ""

    if es == "withdrawn":
        return {"code": "withdrawn", "label": "سحب ملف"}
    if es == "suspended":
        return {"code": "suspended", "label": "إيقاف قيد"}
    return academic_status


def _parse_csv_set(value: str) -> set[str]:
    if not value:
        return set()
    return {x.strip() for x in str(value).split(",") if x and x.strip()}


def _enrollment_label_ar(enrollment_status: str) -> str:
    es = str(enrollment_status or "active").strip().lower()
    return {
        "active": "مسجّل",
        "withdrawn": "سحب ملف",
        "suspended": "إيقاف قيد",
        "graduated": "خريج",
    }.get(es, enrollment_status or "—")


def _format_status_action_period(term: str, year: str) -> str:
    t = (term or "").strip()
    y = (year or "").strip()
    if t and y:
        return f"{t} — {y}"
    return t or y


def _fetch_students_performance_meta(cur) -> list[dict]:
    cols = [row[1] for row in cur.execute("PRAGMA table_info(students)").fetchall()]
    has_es = "enrollment_status" in cols
    parts = [
        "student_id",
        "COALESCE(student_name,'') AS student_name",
        "COALESCE(join_year,'') AS join_year",
    ]
    if has_es:
        parts.append("COALESCE(enrollment_status,'active') AS enrollment_status")
        if "status_changed_at" in cols:
            parts.append("status_changed_at")
        if "status_reason" in cols:
            parts.append("COALESCE(status_reason,'') AS status_reason")
        if "status_changed_term" in cols:
            parts.append("COALESCE(status_changed_term,'') AS status_changed_term")
        if "status_changed_year" in cols:
            parts.append("COALESCE(status_changed_year,'') AS status_changed_year")
    sql = "SELECT " + ", ".join(parts) + " FROM students ORDER BY student_id"
    out: list[dict] = []
    for r in cur.execute(sql):
        d = {
            "student_id": r["student_id"],
            "student_name": r["student_name"] or "",
            "join_year": r["join_year"] or "",
            "enrollment_status": (r["enrollment_status"] if has_es else "active"),
        }
        if has_es:
            if "status_changed_at" in r.keys():
                d["status_changed_at"] = r["status_changed_at"]
            else:
                d["status_changed_at"] = None
            d["status_reason"] = (r["status_reason"] if "status_reason" in r.keys() else "") or ""
            d["status_changed_term"] = (
                (r["status_changed_term"] if "status_changed_term" in r.keys() else "") or ""
            )
            d["status_changed_year"] = (
                (r["status_changed_year"] if "status_changed_year" in r.keys() else "") or ""
            )
        else:
            d["status_changed_at"] = None
            d["status_reason"] = ""
            d["status_changed_term"] = ""
            d["status_changed_year"] = ""
        out.append(d)
    return out


def _build_filter_summary_ar(
    status_codes: set[str],
    enrollment_values: set[str],
    exclude_enrollment: set[str],
    student_ids: set[str],
) -> str:
    parts = []
    status_labels = {
        "good": "سليم أكاديمياً",
        "warning_1": "إنذار أول",
        "warning_2": "إنذار ثانٍ",
        "warning_3": "أكثر من إنذارين",
        "no_data": "لا توجد بيانات درجات",
        "withdrawn": "سحب ملف (حالة قيد)",
        "suspended": "إيقاف قيد (حالة قيد)",
    }
    if status_codes:
        parts.append("الحالة الظاهرة في التقرير: " + "، ".join(sorted(status_labels.get(s, s) for s in status_codes)))
    en_labels = {"active": "مسجّل", "withdrawn": "سحب ملف", "suspended": "إيقاف قيد", "graduated": "خريج"}
    if enrollment_values:
        parts.append("حالة القيد: " + "، ".join(en_labels.get(e, e) for e in sorted(enrollment_values)))
    if exclude_enrollment:
        parts.append("مستثنى من التصدير: " + "، ".join(en_labels.get(e, e) for e in sorted(exclude_enrollment)))
    if student_ids:
        parts.append(f"طلبة محددون: {len(student_ids)}")
    return " — ".join(parts) if parts else "بدون قيود إضافية (جميع الطلبة المطابقة لباقي المعايير)"


def _prepare_rows_for_print(rows: list[dict]) -> list[dict]:
    out = []
    for r in rows:
        item = dict(r)
        item["enrollment_label_ar"] = _enrollment_label_ar(item.get("enrollment_status"))
        item["status_action_period_ar"] = _format_status_action_period(
            item.get("status_changed_term"), item.get("status_changed_year")
        )
        out.append(item)
    return out


def _build_performance_export_rows(conn) -> list[dict]:
    cur = conn.cursor()
    meta_rows = _fetch_students_performance_meta(cur)

    data_rows: list[dict] = []
    for r in meta_rows:
        sid = r["student_id"]
        name = r["student_name"]
        join_year = r["join_year"]
        enrollment_status = r["enrollment_status"]
        status_changed_at = r.get("status_changed_at")
        status_reason = r.get("status_reason") or ""
        status_changed_term = r.get("status_changed_term") or ""
        status_changed_year = r.get("status_changed_year") or ""

        tr = _load_transcript_data(sid)
        ordered = tr.get("ordered_semesters", []) or []
        sem_gpas = tr.get("semester_gpas", {}) or {}
        cumulative_gpa = tr.get("cumulative_gpa", 0.0)
        completed_units = int(tr.get("completed_units") or 0)

        last_semester = ordered[-1] if len(ordered) >= 1 else ""
        prev_semester = ordered[-2] if len(ordered) >= 2 else ""
        third_semester = ordered[-3] if len(ordered) >= 3 else ""

        last_gpa = sem_gpas.get(last_semester, None) if last_semester else None
        prev_gpa = sem_gpas.get(prev_semester, None) if prev_semester else None
        third_gpa = sem_gpas.get(third_semester, None) if third_semester else None

        status = _compute_status(ordered, sem_gpas, cumulative_gpa)
        status = _override_status_by_enrollment(enrollment_status, status)

        exc_row = cur.execute(
            """
            SELECT id, type, note, is_active
            FROM student_exceptions
            WHERE student_id = ? AND type = 'extra_chance'
            ORDER BY id DESC
            LIMIT 1
            """,
            (sid,),
        ).fetchone()
        has_extra = bool(exc_row and exc_row[3])
        extra_note = exc_row[2] if exc_row else ""

        data_rows.append(
            {
                "student_id": sid,
                "student_name": name,
                "join_year": join_year,
                "enrollment_status": enrollment_status,
                "status_changed_at": status_changed_at,
                "status_reason": status_reason,
                "status_changed_term": status_changed_term,
                "status_changed_year": status_changed_year,
                "last_semester": last_semester,
                "last_semester_gpa": last_gpa,
                "prev_semester": prev_semester,
                "prev_semester_gpa": prev_gpa,
                "third_semester": third_semester,
                "third_semester_gpa": third_gpa,
                "cumulative_gpa": cumulative_gpa,
                "completed_units": completed_units,
                "status_code": status["code"],
                "status_label": status["label"],
                "extra_chance": has_extra,
                "extra_chance_note": extra_note or "",
            }
        )
    return data_rows


def _performance_docx_response(rows: list[dict], filter_summary: str):
    import tempfile

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return (
            jsonify({"status": "error", "message": "مكتبة python-docx غير متوفرة. ثبّت docxtpl أو python-docx."}),
            500,
        )

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = t.add_run("جامعة درنة — كلية الهندسة — قسم الهندسة الميكانيكية")
    r0.bold = True

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = t2.add_run("تقرير أداء الطلبة (وثيقة تقريرية)")
    r1.bold = True

    doc.add_paragraph(f"تاريخ الإصدار: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC")
    doc.add_paragraph(f"معايير الاستخراج: {filter_summary}")

    headers = [
        "الرقم الدراسي",
        "الاسم",
        "سنة الالتحاق",
        "حالة القيد",
        "فصل وسنة الإجراء",
        "الفصل الأخير",
        "الفصل السابق",
        "المعدل التراكمي",
        "الوحدات",
        "الحالة / الملاحظات",
        "فرصة استثنائية",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    def _sem_gpa(sem: str, gpa):
        if not sem:
            return "—"
        if gpa is None:
            return f"{sem} (—)"
        try:
            return f"{sem} ({float(gpa):.2f})"
        except Exception:
            return f"{sem} (—)"

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = str(row.get("student_id") or "")
        cells[1].text = str(row.get("student_name") or "")
        cells[2].text = str(row.get("join_year") or "")
        cells[3].text = _enrollment_label_ar(row.get("enrollment_status"))
        sap = _format_status_action_period(
            row.get("status_changed_term"), row.get("status_changed_year")
        )
        cells[4].text = sap or "—"
        cells[5].text = _sem_gpa(row.get("last_semester") or "", row.get("last_semester_gpa"))
        cells[6].text = _sem_gpa(row.get("prev_semester") or "", row.get("prev_semester_gpa"))
        try:
            cg = row.get("cumulative_gpa")
            cells[7].text = f"{float(cg):.2f}" if cg is not None else "—"
        except Exception:
            cells[7].text = "—"
        cells[8].text = str(int(row.get("completed_units") or 0))
        cells[9].text = str(row.get("status_label") or "")
        ex = "لا"
        if row.get("extra_chance"):
            note = (row.get("extra_chance_note") or "").strip()
            ex = "نعم" + (f" — {note}" if note else "")
        cells[10].text = ex

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _fr = foot.add_run("مُنشأ آلياً — لا يُعتمد دون توقيع الجهة المختصة عند الحاجة.")
    _fr.italic = True

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = tmp.name
    tmp.close()
    doc.save(tmp_path)
    fname = f"performance_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(
        tmp_path,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=fname,
    )


def _apply_export_filters(
    rows: list[dict],
    status_codes: set[str],
    enrollment_values: set[str],
    exclude_enrollment: set[str],
    student_ids: set[str],
) -> list[dict]:
    out = []
    for r in rows:
        sid = str(r.get("student_id") or "").strip()
        sc = str(r.get("status_code") or "").strip()
        es = str(r.get("enrollment_status") or "active").strip().lower()

        if status_codes and sc not in status_codes:
            continue
        if enrollment_values and es not in enrollment_values:
            continue
        if exclude_enrollment and es in exclude_enrollment:
            continue
        if student_ids and sid not in student_ids:
            continue
        out.append(r)
    return out


@performance_bp.route("/report")
@role_required("admin", "supervisor")
def performance_report():
    """
    تقرير موجز لأداء الطلبة:
    - آخر فصلين (أو ثلاثة فصول) مع معدلاتهم.
    - المعدل التراكمي والوحدات المنجزة.
    - حالة تقريبية حسب لائحة الإنذار (مادة 43) وبعض إشارات الفصل (مادة 44).
    """
    results = []

    with get_connection() as conn:
        cur = conn.cursor()
        meta_rows = _fetch_students_performance_meta(cur)

        # Scope enforcement:
        # supervisor يرى فقط طلابه المسندين، بينما admin يرى الجميع
        user_role = session.get("user_role")
        is_supervisor = (user_role == "supervisor") or (user_role == "instructor" and int(session.get("is_supervisor") or 0) == 1)
        if is_supervisor or user_role == "instructor":
            instructor_id = session.get("instructor_id")
            if not instructor_id:
                return jsonify({"students": []})

            allowed_student_ids = set()
            if is_supervisor:
                allowed_student_ids = {
                    r[0]
                    for r in cur.execute(
                        "SELECT student_id FROM student_supervisor WHERE instructor_id = ?",
                        (instructor_id,),
                    ).fetchall()
                }
            else:
                instr_row = cur.execute(
                    "SELECT name FROM instructors WHERE id = ? LIMIT 1",
                    (instructor_id,),
                ).fetchone()
                instructor_name = instr_row[0] if instr_row else ""
                if not instructor_name:
                    return jsonify({"students": []})

                term_name, term_year = get_current_term(conn=conn)
                semester_label = f"{(term_name or '').strip()} {(term_year or '').strip()}".strip()
                if not semester_label:
                    return jsonify({"students": []})

                allowed_student_ids = {
                    r[0]
                    for r in cur.execute(
                        """
                        SELECT DISTINCT r.student_id
                        FROM schedule s
                        JOIN registrations r ON r.course_name = s.course_name
                        WHERE s.semester = ?
                          AND s.instructor = ?
                          AND COALESCE(s.course_name,'') <> ''
                        """,
                        (semester_label, instructor_name),
                    ).fetchall()
                }

            if allowed_student_ids:
                meta_rows = [r for r in meta_rows if r.get("student_id") in allowed_student_ids]
            else:
                meta_rows = []

        for r in meta_rows:
            sid = r["student_id"]
            name = r["student_name"]
            join_year = r["join_year"]
            enrollment_status = r["enrollment_status"]
            status_changed_at = r.get("status_changed_at")
            status_reason = r.get("status_reason") or ""
            status_changed_term = r.get("status_changed_term") or ""
            status_changed_year = r.get("status_changed_year") or ""

            data = _load_transcript_data(sid)
            ordered = data.get("ordered_semesters", []) or []
            sem_gpas = data.get("semester_gpas", {}) or {}
            cumulative_gpa = data.get("cumulative_gpa", 0.0)
            completed_units = int(data.get("completed_units") or 0)

            last_semester = ordered[-1] if len(ordered) >= 1 else ""
            prev_semester = ordered[-2] if len(ordered) >= 2 else ""
            third_semester = ordered[-3] if len(ordered) >= 3 else ""

            last_gpa = sem_gpas.get(last_semester, None) if last_semester else None
            prev_gpa = sem_gpas.get(prev_semester, None) if prev_semester else None
            third_gpa = sem_gpas.get(third_semester, None) if third_semester else None

            status = _compute_status(ordered, sem_gpas, cumulative_gpa)
            status = _override_status_by_enrollment(enrollment_status, status)

            exc_row = cur.execute(
                """
                SELECT id, type, note, is_active
                FROM student_exceptions
                WHERE student_id = ? AND type = 'extra_chance'
                ORDER BY id DESC
                LIMIT 1
                """,
                (sid,),
            ).fetchone()
            has_extra = bool(exc_row and exc_row[3])
            extra_note = exc_row[2] if exc_row else ""

            results.append(
                {
                    "student_id": sid,
                    "student_name": name,
                    "join_year": join_year,
                    "enrollment_status": enrollment_status,
                    "status_changed_at": status_changed_at,
                    "status_reason": status_reason,
                    "status_changed_term": status_changed_term,
                    "status_changed_year": status_changed_year,
                    "last_semester": last_semester,
                    "last_semester_gpa": last_gpa,
                    "prev_semester": prev_semester,
                    "prev_semester_gpa": prev_gpa,
                    "third_semester": third_semester,
                    "third_semester_gpa": third_gpa,
                    "cumulative_gpa": cumulative_gpa,
                    "completed_units": completed_units,
                    "status_code": status["code"],
                    "status_label": status["label"],
                    "extra_chance": has_extra,
                    "extra_chance_note": extra_note or "",
                }
            )

    return jsonify({"students": results})


@performance_bp.route("/status/<student_id>")
@role_required("admin", "supervisor", "student")
def performance_status(student_id: str):
    """
    إرجاع حالة أكاديمية موجزة لطالب واحد:
    - status_code, status_label
    - cumulative_gpa, completed_units
    - extra_chance (فرصة استثنائية) إن وُجدت.
    يستخدمها كشف الدرجات لعرض ملاحظة سريعة.
    """
    sid = (student_id or "").strip()
    if not sid:
        return jsonify({"status": "error", "message": "student_id مطلوب"}), 400

    with get_connection() as conn:
        cur = conn.cursor()

        # Scope enforcement:
        # student يرى فقط سجله، supervisor يرى فقط طلابه
        user_role = session.get("user_role")
        if user_role == "student":
            sid_session = session.get("student_id") or session.get("user")
            if (sid_session or "").strip() != sid:
                return jsonify({"status": "error", "message": "FORBIDDEN"}), 403

        is_supervisor = (user_role == "supervisor") or (user_role == "instructor" and int(session.get("is_supervisor") or 0) == 1)
        if is_supervisor:
            instructor_id = session.get("instructor_id")
            if not instructor_id:
                return jsonify({"status": "error", "message": "FORBIDDEN"}), 403
            allowed = cur.execute(
                """
                SELECT 1
                FROM student_supervisor
                WHERE student_id = ? AND instructor_id = ?
                LIMIT 1
                """,
                (sid, instructor_id),
            ).fetchone()
            if not allowed:
                return jsonify({"status": "error", "message": "FORBIDDEN"}), 403
        if user_role == "instructor":
            instructor_id = session.get("instructor_id")
            if not instructor_id:
                return jsonify({"status": "error", "message": "FORBIDDEN"}), 403
            instr_row = cur.execute(
                "SELECT name FROM instructors WHERE id = ? LIMIT 1",
                (instructor_id,),
            ).fetchone()
            instructor_name = instr_row[0] if instr_row else ""
            if not instructor_name:
                return jsonify({"status": "error", "message": "FORBIDDEN"}), 403

            term_name, term_year = get_current_term(conn=conn)
            semester_label = f"{(term_name or '').strip()} {(term_year or '').strip()}".strip()
            if not semester_label:
                return jsonify({"status": "error", "message": "FORBIDDEN"}), 403

            allowed = cur.execute(
                """
                SELECT 1
                FROM schedule s
                JOIN registrations r ON r.course_name = s.course_name
                WHERE s.semester = ?
                  AND s.instructor = ?
                  AND r.student_id = ?
                LIMIT 1
                """,
                (semester_label, instructor_name, sid),
            ).fetchone()
            if not allowed:
                return jsonify({"status": "error", "message": "FORBIDDEN"}), 403

        cols = [row[1] for row in cur.execute("PRAGMA table_info(students)").fetchall()]
        has_enrollment_status = "enrollment_status" in cols

        data = _load_transcript_data(sid)
        ordered = data.get("ordered_semesters", []) or []
        sem_gpas = data.get("semester_gpas", {}) or {}
        cumulative_gpa = data.get("cumulative_gpa", 0.0)
        completed_units = int(data.get("completed_units") or 0)

        status = _compute_status(ordered, sem_gpas, cumulative_gpa)

        enrollment_status = "active"
        if has_enrollment_status:
            row = cur.execute(
                "SELECT COALESCE(enrollment_status,'active') FROM students WHERE student_id = ?",
                (sid,),
            ).fetchone()
            enrollment_status = row[0] if row else "active"
        status = _override_status_by_enrollment(enrollment_status, status)

        exc_row = cur.execute(
            """
            SELECT id, type, note, is_active
            FROM student_exceptions
            WHERE student_id = ? AND type = 'extra_chance'
            ORDER BY id DESC
            LIMIT 1
            """,
            (sid,),
        ).fetchone()
        has_extra = bool(exc_row and exc_row[3])
        extra_note = exc_row[2] if exc_row else ""

    return jsonify(
        {
            "student_id": sid,
            "status_code": status["code"],
            "status_label": status["label"],
            "cumulative_gpa": cumulative_gpa,
            "completed_units": completed_units,
            "extra_chance": has_extra,
            "extra_chance_note": extra_note or "",
        }
    )


@performance_bp.route("/extra_chance", methods=["POST"])
@role_required("admin")
def set_extra_chance():
    """
    منح / إلغاء فرصة استثنائية لطالب.
    body:
      - student_id (مطلوب)
      - active: true/false (مطلوب)
      - note: ملاحظة اختيارية
      - created_by: اسم المستخدم (اختياري؛ يمكن إرساله من الواجهة)
    """
    data = request.get_json(force=True) or {}
    sid = (data.get("student_id") or "").strip()
    active_raw = data.get("active")
    note = (data.get("note") or "").strip() or None
    created_by = (data.get("created_by") or "").strip() or None

    if not sid:
        return jsonify({"status": "error", "message": "student_id مطلوب"}), 400

    is_active = 1 if bool(active_raw) else 0

    from datetime import datetime

    now = datetime.utcnow().isoformat()

    with get_connection() as conn:
        cur = conn.cursor()
        cur.execute(
            """
            INSERT INTO student_exceptions (student_id, type, note, created_by, created_at, is_active)
            VALUES (?, 'extra_chance', ?, ?, ?, ?)
            """,
            (sid, note, created_by, now, is_active),
        )
        conn.commit()

    return jsonify({"status": "ok"})


@performance_bp.route("/export")
@role_required("admin")
def export_performance():
    """
    تصدير تقرير الأداء مع نفس فلاتر الاستعلام:
      - format=xlsx (افتراضي) | pdf | docx
      - status_codes, enrollment, exclude_enrollment, student_ids
    """
    from pandas import DataFrame

    fmt = (request.args.get("format") or "xlsx").lower().strip()
    status_codes = _parse_csv_set(request.args.get("status_codes", ""))
    enrollment_values = _parse_csv_set(request.args.get("enrollment", ""))
    exclude_enrollment = _parse_csv_set(request.args.get("exclude_enrollment", ""))
    student_ids = _parse_csv_set(request.args.get("student_ids", ""))

    with get_connection() as conn:
        data_rows = _build_performance_export_rows(conn)

    data_rows = _apply_export_filters(
        data_rows,
        status_codes=status_codes,
        enrollment_values=enrollment_values,
        exclude_enrollment=exclude_enrollment,
        student_ids=student_ids,
    )

    summary = _build_filter_summary_ar(
        status_codes, enrollment_values, exclude_enrollment, student_ids
    )
    generated_at = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")

    if fmt in ("pdf",):
        print_rows = _prepare_rows_for_print(data_rows)
        html = render_template(
            "performance_report_export_print.html",
            rows=print_rows,
            filter_summary=summary,
            generated_at=generated_at,
        )
        return pdf_response_from_html(html, filename_prefix="performance_report")

    if fmt in ("docx", "word"):
        return _performance_docx_response(data_rows, summary)

    # Excel (افتراضي)
    excel_rows = []
    for r in data_rows:
        er = dict(r)
        er["حالة القيد"] = _enrollment_label_ar(er.get("enrollment_status"))
        er["فصل وسنة الإجراء"] = _format_status_action_period(
            er.get("status_changed_term"), er.get("status_changed_year")
        ) or "—"
        er["ملاحظة القيد"] = (er.get("status_reason") or "").strip()
        er.pop("enrollment_status", None)
        er.pop("status_code", None)
        er.pop("third_semester", None)
        er.pop("third_semester_gpa", None)
        er.pop("status_changed_at", None)
        er.pop("status_reason", None)
        er.pop("status_changed_term", None)
        er.pop("status_changed_year", None)
        excel_rows.append(er)

    preferred = [
        "student_id",
        "student_name",
        "join_year",
        "حالة القيد",
        "فصل وسنة الإجراء",
        "ملاحظة القيد",
        "last_semester",
        "last_semester_gpa",
        "prev_semester",
        "prev_semester_gpa",
        "cumulative_gpa",
        "completed_units",
        "status_label",
        "extra_chance",
        "extra_chance_note",
    ]
    df = DataFrame(excel_rows)
    df = df[[c for c in preferred if c in df.columns]]
    return excel_response_from_df(df, filename_prefix="performance_report")

