"""أدوات دعم التصدير العربي: RTL، خطوط، وتنسيق Excel/Word/PDF."""

from __future__ import annotations

from typing import Any

import pandas as pd

ARABIC_FONT_PRIMARY = "Arial"
ARABIC_FONT_FALLBACK = "Tahoma"
PDF_ARABIC_FONT_STACK = "Tahoma, Arial, 'DejaVu Sans', sans-serif"
PDF_ARABIC_FONT_STACK_WEB = "'Cairo', 'Noto Naskh Arabic', Tahoma, Arial, sans-serif"

WKHTMLTOPDF_ARABIC_OPTIONS: dict[str, str | None] = {
    "encoding": "UTF-8",
    "page-size": "A4",
    "margin-top": "14mm",
    "margin-right": "12mm",
    "margin-bottom": "14mm",
    "margin-left": "12mm",
    "disable-smart-shrinking": None,
    "enable-local-file-access": None,
    "print-media-type": None,
}


def pdf_arabic_extra_css(*, for_pdf: bool = False) -> str:
    """CSS إضافي لقوالب PDF/HTML العربية."""
    font_stack = PDF_ARABIC_FONT_STACK if for_pdf else PDF_ARABIC_FONT_STACK_WEB
    font_import = ""
    if not for_pdf:
        font_import = "@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700&display=swap');"
    pdf_table_rules = ""
    page_rules = ""
    if for_pdf:
        page_rules = """
    @page {
      size: A4;
      margin: 14mm 12mm 16mm 12mm;
    }
    """
        pdf_table_rules = """
    * { box-sizing: border-box; }
    html, body {
      max-width: 100%;
      overflow-x: hidden;
      padding: 0 !important;
      margin: 0 !important;
    }
    h1, h2, h3, h4, p, li, blockquote, td, th {
      overflow-wrap: anywhere;
      word-wrap: break-word;
      word-break: break-word;
      hyphens: auto;
    }
    table {
      table-layout: fixed;
      width: 100% !important;
      max-width: 100% !important;
      border-collapse: collapse;
    }
    th, td {
      border: 1px solid #666 !important;
      padding: 4px 6px !important;
      line-height: 1.4;
      max-width: 0; /* مع table-layout:fixed يفرض التفاف النص */
    }
    table.tbl-rec td:nth-child(1) { width: 52%; }
    table.tbl-rec td:nth-child(2) { width: 22%; }
    table.tbl-rec td:nth-child(3) { width: 26%; }
    table.tbl-data { font-size: 9px; }
    table.tbl-data th, table.tbl-data td { padding: 4px 5px; line-height: 1.35; }
    img, canvas { max-width: 100% !important; height: auto !important; }
    .narrative-block { margin: 0 0 10px 0; line-height: 1.55; }
    .summary-panel, .analysis-panel, .analysis-box, .narrative {
      border: 1px solid #c8d6e5;
      background: #f8fafc;
      padding: 10px 12px;
      margin: 10px 0 14px;
      max-width: 100%;
      overflow-wrap: anywhere;
    }
    .summary-panel p, .analysis-panel p, .analysis-panel .narrative-block {
      margin: 0 0 8px 0;
      padding: 0;
      display: block;
      clear: both;
    }
    .open-text-quotes blockquote {
      max-width: 100%;
      overflow-wrap: anywhere;
    }
    .section-block { page-break-inside: avoid; }
    """
    else:
        page_rules = """
    @page {
      size: A4;
      margin: 14mm 12mm 16mm 12mm;
    }
    @media print {
      html, body {
        max-width: 100% !important;
        overflow-x: hidden !important;
        padding: 0 !important;
        margin: 0 !important;
      }
      table {
        table-layout: fixed !important;
        width: 100% !important;
        max-width: 100% !important;
      }
      th, td {
        border: 1px solid #666 !important;
        overflow-wrap: anywhere !important;
        word-break: break-word !important;
        max-width: 0;
      }
      h1, h2, h3, h4, p, li, blockquote {
        overflow-wrap: anywhere !important;
        word-break: break-word !important;
      }
      .chart-box { height: auto !important; min-height: 220px; page-break-inside: avoid; }
      .no-print { display: none !important; }
    }
    """
    return f"""
    {font_import}
    {page_rules}
    html, body {{
      direction: rtl;
      unicode-bidi: embed;
      font-family: {font_stack};
      text-align: right;
    }}
    table {{ direction: rtl; border-collapse: collapse; }}
    th, td {{
      text-align: right;
      vertical-align: top;
      word-wrap: break-word;
      overflow-wrap: break-word;
    }}
    {pdf_table_rules}
    """


def excel_arabic_workbook_formats(workbook) -> dict[str, Any]:
    """تنسيقات Excel موحّدة للعربية."""
    base = {"font_name": ARABIC_FONT_PRIMARY, "align": "right", "valign": "vcenter"}
    return {
        "cover": workbook.add_format({**base, "bold": True, "font_size": 14}),
        "dept": workbook.add_format(
            {**base, "bold": True, "font_size": 12, "font_color": "#0d3b66"}
        ),
        "header": workbook.add_format(
            {**base, "bold": True, "bg_color": "#e8eef4", "border": 1}
        ),
        "cell": workbook.add_format({**base, "border": 1}),
        "wrap": workbook.add_format({**base, "border": 1, "text_wrap": True}),
        "title": workbook.add_format({**base, "bold": True, "font_size": 12, "bg_color": "#f0f6fc"}),
    }


def write_excel_sheet_rtl(
    worksheet,
    data: pd.DataFrame,
    *,
    formats: dict[str, Any],
    cover_sheet: bool = False,
) -> None:
    """كتابة ورقة Excel بمحاذاة عربية واتجاه RTL."""
    worksheet.right_to_left()
    if data.empty:
        return
    cols = list(data.columns)
    if cover_sheet and len(cols) >= 2:
        worksheet.set_column(0, 0, 34)
        worksheet.set_column(1, 1, 52)
        for row_idx, row in data.iterrows():
            label = str(row.iloc[0] if len(row) else "")
            value = str(row.iloc[1] if len(row) > 1 else "")
            if row_idx == 0:
                fmt = formats["cover"]
            elif label == "القسم":
                fmt = formats["dept"]
            else:
                fmt = formats["cell"]
            worksheet.write(int(row_idx) + 1, 0, label, fmt)
            worksheet.write(int(row_idx) + 1, 1, value, fmt)
        return
    for col_idx, col_name in enumerate(cols):
        worksheet.write(0, col_idx, str(col_name), formats["header"])
        width = max(12, min(36, len(str(col_name)) + 6))
        worksheet.set_column(col_idx, col_idx, width)
    wrap_cols = {"الفقرة", "التوصية", "الاستنتاج", "أسباب_الفجوة", "أسباب_الفجوة"}
    for row_idx, row in data.iterrows():
        for col_idx, col_name in enumerate(cols):
            val = row.iloc[col_idx]
            txt = "" if val is None or (isinstance(val, float) and pd.isna(val)) else str(val)
            fmt = formats["wrap"] if str(col_name) in wrap_cols else formats["cell"]
            worksheet.write(int(row_idx) + 1, col_idx, txt, fmt)


def _docx_oxml():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    return OxmlElement, qn


def set_docx_paragraph_rtl(paragraph, *, align_right: bool = True) -> None:
    """ضبط اتجاه فقرة Word إلى RTL."""
    OxmlElement, qn = _docx_oxml()
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)
    if align_right:
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "right")
        p_pr.append(jc)


def set_docx_run_arabic_font(run, font_name: str = ARABIC_FONT_PRIMARY) -> None:
    """خط عربي مع دعم النص المعقّد في Word."""
    OxmlElement, qn = _docx_oxml()
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    r_pr.append(r_fonts)


def docx_add_rtl_paragraph(
    doc,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    center: bool = False,
    font_size: int | None = None,
):
    """إضافة فقرة عربية منسّقة."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_docx_paragraph_rtl(p, align_right=not center)
    run = p.add_run(str(text or ""))
    run.bold = bold
    run.italic = italic
    if font_size:
        run.font.size = Pt(font_size)
    set_docx_run_arabic_font(run)
    return p


def docx_add_rtl_heading(doc, text: str, level: int = 2):
    """عنوان عربي RTL."""
    h = doc.add_heading(str(text or ""), level=level)
    set_docx_paragraph_rtl(h)
    for run in h.runs:
        set_docx_run_arabic_font(run)
    return h


def docx_apply_table_rtl(table) -> None:
    """ضبط جدول Word للعرض من اليمين."""
    OxmlElement, qn = _docx_oxml()
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    bidi_visual = OxmlElement("w:bidiVisual")
    bidi_visual.set(qn("w:val"), "1")
    tbl_pr.append(bidi_visual)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_docx_paragraph_rtl(p)


def docx_fill_rtl_table(table, headers: list[str], rows: list[list[str]]) -> None:
    """ملء جدول Word بعناوين وصفوف عربية."""
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)
    for row_vals in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_vals):
            cells[i].text = str(val if val is not None else "—")
    docx_apply_table_rtl(table)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_docx_run_arabic_font(run)
