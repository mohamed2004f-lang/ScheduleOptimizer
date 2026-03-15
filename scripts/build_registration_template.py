"""
إنشاء قالب استمارة التسجيل registration_form_template.docx في frontend/templates/
يشغّل مرة واحدة: python scripts/build_registration_template.py
"""
import os
import sys

# جذر المشروع
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
sys.path.insert(0, ROOT)

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = os.path.join(ROOT, "frontend", "templates", "registration_form_template.docx")


def main():
    doc = Document()
    doc.add_paragraph()
    # عنوان
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("جامعة درنة — كلية الهندسة")
    r.bold = True
    r.font.size = Pt(14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("قسم {{ department }}")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("استمارة تسجيل مقررات دراسية")
    r.bold = True
    r.font.size = Pt(12)
    doc.add_paragraph()

    # بيانات الطالب (نص بسيط مع متغيرات)
    doc.add_paragraph("اسم الطالب: {{ student_name }}    الرقم الدراسي: {{ student_id }}")
    doc.add_paragraph("الرقم الجامعي: {{ university_number }}    الفصل: {{ semester }}")
    doc.add_paragraph("الوحدات المنجزة: {{ completed_units }}    المعدل التراكمي: {{ cumulative_gpa }}    الحالة: {{ status }}    مجموع وحدات الفصل: {{ total_units }}")
    doc.add_paragraph()

    # جدول المقررات: ترتيب الأعمدة من اليمين لليسار (عند الفتح) = الترقيم، المقرر، الرمز، الوحدات، الملاحظات
    # نضع الملاحظات في العمود الأول (يسار) والترقيم في الأخير (يمين) ليتضح الترتيب صحيحاً عند القراءة من اليمين
    table = doc.add_table(rows=11, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "ملاحظات"
    hdr[1].text = "الوحدات"
    hdr[2].text = "الرمز"
    hdr[3].text = "اسم المقرر"
    hdr[4].text = "م"
    for i in range(10):
        row = table.rows[i + 1].cells
        row[0].text = "{{ courses[" + str(i) + "].notes }}"
        row[1].text = "{{ courses[" + str(i) + "].units }}"
        row[2].text = "{{ courses[" + str(i) + "].code }}"
        row[3].text = "{{ courses[" + str(i) + "].name }}"
        row[4].text = "{{ courses[" + str(i) + "].index }}"

    doc.add_paragraph()
    doc.add_paragraph("توقيع الطالب: _______________     المرشد الأكاديمي: _______________     رئيس القسم: _______________")

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print("تم إنشاء القالب:", OUTPUT_PATH)


if __name__ == "__main__":
    main()
